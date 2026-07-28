from copy import deepcopy
from pathlib import Path

from docx import Document


PAPER = Path(
    "/Users/chengyizhou/Documents/GitHub/solar-pv-prediction/"
    "Solar_PV_Research_Paper_Model_C_Revised.docx"
)


def unique_paragraph(document, prefix):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(prefix)
    ]
    if len(matches) != 1:
        raise ValueError(
            f"Expected one paragraph beginning {prefix!r}; found {len(matches)}"
        )
    return matches[0]


def replace_text(paragraph, text):
    template = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    run = paragraph.add_run(text)
    if template is not None and template._r.rPr is not None:
        run._r.insert(0, deepcopy(template._r.rPr))


def insert_heading_before(document, target_paragraph, text):
    heading = document.add_paragraph(text, style="Heading 2")
    target_paragraph._p.addprevious(heading._p)


def main():
    document = Document(PAPER)

    first_secondary = unique_paragraph(
        document,
        "6.1 Physics baseline substantially improves",
    )
    first_primary = unique_paragraph(
        document,
        "6.5 The primary future-time evaluation",
    )
    discussion = unique_paragraph(document, "7. Discussion")

    body = document._element.body
    children = list(body)
    primary_start = children.index(first_primary._p)
    discussion_start = children.index(discussion._p)
    primary_block = children[primary_start:discussion_start]

    for element in primary_block:
        body.remove(element)
    for element in primary_block:
        first_secondary._p.addprevious(element)

    replace_text(
        first_primary,
        "6.1 Primary future-time evaluation with an explicit time control",
    )
    primary_gain = unique_paragraph(
        document,
        "6.6 Physics adds a modest improvement",
    )
    replace_text(
        primary_gain,
        "6.2 Physics adds a modest improvement beyond ordinary time controls",
    )

    cold_start = unique_paragraph(
        document,
        "The four systems with no eligible pre-2023 history",
    )
    insert_heading_before(
        document,
        cold_start,
        "6.3 No-history systems support the bounded cold-start result",
    )
    bias = unique_paragraph(document, "Bias remains the main weakness")
    insert_heading_before(
        document,
        bias,
        "6.4 Bias remains unresolved beyond the time-aware control",
    )

    replace_text(
        first_secondary,
        "6.5 Secondary same-period physics baseline",
    )
    replace_text(
        unique_paragraph(
            document,
            "6.2 Same-period feature bundles",
        ),
        "6.6 Secondary same-period feature bundles reduce average XGBoost error",
    )
    replace_text(
        unique_paragraph(
            document,
            "6.3 The gain is broad",
        ),
        "6.7 Same-period gains are broad but not universal across systems",
    )
    replace_text(
        unique_paragraph(
            document,
            "6.4 Paired uncertainty",
        ),
        "6.8 Same-period paired uncertainty supports the within-campus comparison",
    )

    replace_text(
        unique_paragraph(document, "Table 5. Primary future-time"),
        "Table 4. Primary future-time performance across 37 held-out PV systems",
    )
    replace_text(
        unique_paragraph(document, "Table 4. Macro performance"),
        "Table 5. Secondary same-period performance across 37 held-out PV systems",
    )
    replace_text(
        unique_paragraph(document, "Figure 3. Primary future-time"),
        "Figure 2. Primary future-time Model C comparison. The left panel shows "
        "macro nRMSE for all five models when training ends in 2022 and testing "
        "uses held-out systems in 2023. The right panel compares physics-enhanced "
        "Model B with weather-plus-time Model C; Model B is lower on 34 of 37 "
        "systems. Open circles identify four systems with no eligible pre-2023 "
        "history. Source: Model C future-time system-level results.",
    )
    replace_text(
        unique_paragraph(document, "Figure 2. Secondary same-period"),
        "Figure 3. Secondary same-period feature-bundle comparison. The left "
        "panel shows macro nRMSE across four models. The right panel compares "
        "weather-only Model A with physics-enhanced Model B for each held-out "
        "system; 36 of 37 points favor Model B. Because training systems supply "
        "labels for the same timestamps, the figure does not separate physics "
        "from ordinary time context. Source: Study analysis outputs.",
    )

    document.save(PAPER)
    print(f"Reordered primary results in {PAPER}")


if __name__ == "__main__":
    main()
