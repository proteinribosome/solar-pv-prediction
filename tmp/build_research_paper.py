from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "tmp" / "paper_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = ROOT / "Solar_PV_Research_Paper_Draft.docx"

NAVY = "1F4E79"
DEEP_NAVY = "17365D"
BLUE = "5B9BD5"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF4F9"
GOLD = "C69214"
PALE_GOLD = "FFF4CE"
INK = "202124"
MUTED = "5F6368"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E1E8"
WHITE = "FFFFFF"
RED = "9C0006"


def set_run_font(run, name="Calibri", size=11, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths_dxa, *, numeric_cols=(), caption=None, source=None):
    if caption:
        p = doc.add_paragraph(style="Caption")
        p.paragraph_format.keep_with_next = True
        r = p.add_run(caption)
        set_run_font(r, size=9.5, bold=True, color=DEEP_NAVY)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=9.2,
                      align=WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT)
        shade_cell(table.rows[0].cells[idx], NAVY)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align)
            if ridx % 2 == 1:
                shade_cell(cells[idx], LIGHT_GRAY)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    if source:
        p = doc.add_paragraph(style="Table Source")
        p.paragraph_format.keep_with_next = False
        r = p.add_run(source)
        set_run_font(r, size=8.5, italic=True, color=MUTED)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.keep_with_next = False
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=False, italic=False, color=INK)
    return p


def add_body(doc, text, *, bold_lead=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, name="Cambria Math", size=11.5, italic=True, color=INK)
    return p


def add_note_box(doc, title, body, *, fill=PALE_GOLD, title_color=RED):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    shade_cell(cell, fill)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=GOLD, size=8)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=INK)
    return table


def make_figures():
    font_dir = Path("/System/Library/Fonts/Supplemental")
    regular = lambda s: ImageFont.truetype(str(font_dir / "Arial.ttf"), s)
    bold = lambda s: ImageFont.truetype(str(font_dir / "Arial Bold.ttf"), s)

    def center_multiline(draw, box, text, font, fill, spacing=8):
        x0, y0, x1, y1 = box
        bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.multiline_text(((x0 + x1 - tw) / 2, (y0 + y1 - th) / 2), text,
                            font=font, fill=fill, spacing=spacing, align="center")

    def arrow(draw, start, end, fill="#5F6368", width=5):
        draw.line([start, end], fill=fill, width=width)
        dx, dy = end[0] - start[0], end[1] - start[1]
        length = max((dx * dx + dy * dy) ** 0.5, 1)
        ux, uy = dx / length, dy / length
        px, py = -uy, ux
        tip = end
        left = (end[0] - 22 * ux + 10 * px, end[1] - 22 * uy + 10 * py)
        right = (end[0] - 22 * ux - 10 * px, end[1] - 22 * uy - 10 * py)
        draw.polygon([tip, left, right], fill=fill)

    # Figure 1: experimental design schematic.
    fig = Image.new("RGB", (2200, 1120), "white")
    draw = ImageDraw.Draw(fig)
    draw.text((110, 70), "Experimental design for cross-system prediction", font=bold(50), fill=f"#{DEEP_NAVY}")
    draw.text((110, 135), "All observations from a held-out system remain outside training.", font=regular(30), fill=f"#{MUTED}")
    boxes = [
        ((90, 260, 500, 455), "37 rooftop PV systems\n2021-2023", PALE_BLUE, NAVY),
        ((600, 260, 1010, 455), "Daylight + quality control\n1.37M ready rows", PALE_BLUE, NAVY),
        ((1110, 260, 1520, 455), "Five grouped folds\n7-8 systems held out", PALE_BLUE, NAVY),
        ((1620, 260, 2080, 455), "System-level metrics\nnRMSE, MAE, MBE", PALE_BLUE, NAVY),
        ((460, 650, 930, 850), "Model A\n4 weather features", PALE_BLUE, NAVY),
        ((1270, 650, 1810, 850), "Model B\n4 weather + 5 physics features", "FFF2CC", GOLD),
    ]
    for box, label, fc, ec in boxes:
        draw.rounded_rectangle(box, radius=18, fill=f"#{fc}", outline=f"#{ec}", width=5)
        center_multiline(draw, box, label, bold(31) if "Model" in label else regular(31), f"#{INK}")
    arrow(draw, (510, 358), (585, 358))
    arrow(draw, (1020, 358), (1095, 358))
    arrow(draw, (1530, 358), (1605, 358))
    arrow(draw, (1280, 470), (820, 635))
    arrow(draw, (1370, 470), (1510, 635))
    arrow(draw, (700, 865), (1030, 970))
    arrow(draw, (1510, 865), (1170, 970))
    draw.rounded_rectangle((640, 960, 1570, 1065), radius=18, fill="#F2F4F7", outline=f"#{MUTED}", width=3)
    center_multiline(draw, (640, 960, 1570, 1065),
                     "Paired comparison isolates the value of physics-derived features",
                     bold(29), f"#{DEEP_NAVY}")
    path1 = ASSET_DIR / "figure_1_experimental_design.png"
    fig.save(path1, dpi=(220, 220))

    # Figure 2: four-model macro nRMSE comparison.
    labels = ["Training mean", "GHI physics", "Weather-only\nXGBoost", "Physics-enhanced\nXGBoost"]
    values = [0.2160, 0.1237, 0.0919, 0.0790]
    colors = ["#D9E1E8", "#A9C4DE", "#5B9BD5", "#C69214"]
    edges = ["#7F8C8D", "#1F4E79", "#1F4E79", "#8A6500"]
    fig = Image.new("RGB", (2100, 1200), "white")
    draw = ImageDraw.Draw(fig)
    draw.text((120, 70), "Cross-system prediction error by model", font=bold(50), fill=f"#{DEEP_NAVY}")
    draw.text((120, 140), "Five-fold leave-systems-out evaluation; n = 37 systems; lower is better",
              font=regular(30), fill=f"#{MUTED}")
    plot_left, plot_right, plot_top, plot_bottom = 650, 1940, 260, 1010
    max_x = 0.24
    for tick in [0.00, 0.05, 0.10, 0.15, 0.20]:
        x = plot_left + (plot_right - plot_left) * tick / max_x
        draw.line([(x, plot_top), (x, plot_bottom)], fill="#E5E7EB", width=3)
        txt = f"{tick:.2f}"
        bbox = draw.textbbox((0, 0), txt, font=regular(26))
        draw.text((x - (bbox[2]-bbox[0])/2, plot_bottom + 18), txt, font=regular(26), fill=f"#{MUTED}")
    bar_h = 115
    row_gap = 55
    for idx, (label, value, color, edge) in enumerate(zip(labels, values, colors, edges)):
        y0 = plot_top + idx * (bar_h + row_gap)
        y1 = y0 + bar_h
        label_box = (90, y0 - 5, 610, y1 + 5)
        center_multiline(draw, label_box, label, regular(31), f"#{INK}", spacing=4)
        x1 = plot_left + (plot_right - plot_left) * value / max_x
        draw.rounded_rectangle((plot_left, y0, x1, y1), radius=12, fill=color, outline=edge, width=4)
        draw.text((x1 + 18, y0 + 34), f"{value:.4f}", font=bold(31), fill=f"#{INK}")
    axis_label = "Macro nRMSE (fraction of installed capacity)"
    bbox = draw.textbbox((0, 0), axis_label, font=regular(29))
    draw.text(((plot_left + plot_right - (bbox[2]-bbox[0]))/2, 1100), axis_label,
              font=regular(29), fill=f"#{INK}")
    arrow(draw, (1740, 790), (1090, 865), fill="#7A5A00", width=5)
    draw.text((1510, 700), "14.0% lower than\nweather-only model", font=bold(27),
              fill="#7A5A00", spacing=4)
    path2 = ASSET_DIR / "figure_2_model_nrmse.png"
    fig.save(path2, dpi=(240, 240))
    return path1, path2


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, NAVY, 12, 6),
        ("Heading 3", 12, DEEP_NAVY, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.208

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.line_spacing = 1.05
    caption.paragraph_format.keep_with_next = True

    if "Table Source" not in styles:
        st = styles.add_style("Table Source", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles["Table Source"]
    st.font.name = "Calibri"
    st.font.size = Pt(8.5)
    st.font.italic = True
    st.font.color.rgb = RGBColor.from_string(MUTED)
    st.paragraph_format.space_before = Pt(4)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.0

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("PHYSICS FEATURES FOR CROSS-SYSTEM PV PREDICTION")
    set_run_font(hr, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])

    doc.core_properties.title = "Physics-Derived Features and Cross-System Generalization in Rooftop Solar PV Power Prediction"
    doc.core_properties.subject = "Pioneer Research Program paper draft"
    doc.core_properties.author = "Student author - name to be added"
    doc.core_properties.keywords = "solar PV, physics-informed machine learning, XGBoost, domain generalization"


def title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Physics-Derived Features and Cross-System Generalization in Rooftop Solar PV Power Prediction")
    set_run_font(r, size=24, bold=True, color=DEEP_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    r = p.add_run("A controlled leave-systems-out evaluation on 37 Hong Kong rooftop PV systems")
    set_run_font(r, size=13.5, italic=True, color=MUTED)

    for label, value in (
        ("Author", "[Student name]"),
        ("Research concentration", "Data Science Using Machine Learning"),
        ("Research mentor", "Professor Sanjay Ranka"),
        ("Program", "Pioneer Research Program, Spring 2026"),
        ("Date", "[Submission date]"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        rl = p.add_run(f"{label}: ")
        set_run_font(rl, size=11, bold=True, color=INK)
        rv = p.add_run(value)
        set_run_font(rv, size=11, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESEARCH DRAFT - NOT READY FOR SUBMISSION")
    set_run_font(r, size=10, bold=True, color=RED)
    doc.add_page_break()


def draft_status_page(doc):
    doc.add_heading("Draft status and academic-integrity note", level=1)
    add_note_box(
        doc,
        "Remove this page before final submission",
        "This document is an AI-assisted initial draft assembled from the project's notebook, saved Phase 3 results, the course syllabus, and cited primary literature. The student must verify every statement, rerun the final analysis, revise the prose into the student's own scholarly voice, and follow the program's required method for citing AI-generated text. The syllabus explicitly requires citation of AI-generated text even when the wording is changed.",
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)
    spacer.paragraph_format.line_spacing = 1.0
    add_body(doc, "The following items are intentionally unfinished because the available repository does not yet contain the saved Phase 4 and Phase 5 system-level result files needed to verify them:")
    for item in (
        "Run Phase 5 and save phase4_system_results.csv, phase5_system_results.csv, phase5_headline.csv, and the paired bootstrap confidence interval.",
        "Replace every bracketed [TO COMPLETE] marker, especially the confidence interval, number of systems improved by Model B, repository URL, author name, and exact software versions.",
        "Add the system-level Model A versus Model B scatterplot after the paired results are saved; this is needed to show whether the 14.0% aggregate improvement is broad or driven by a subset of systems.",
        "Complete the planned SMA robustness analysis or clearly state that it was not performed.",
        "Check all citations against the original articles and the program's required APA guidance, obtain formal topic approval, and incorporate mentor, peer, and Writing Center feedback.",
        "Run Turnitin, resolve genuine quotation or attribution problems, and ensure the final paper represents the student's own analysis and writing.",
    ):
        add_bullet(doc, item)
    add_body(doc, "Evidence status. Phase 3 values are available as CSV files in the repository. Phase 4 values are present as executed notebook outputs but are not currently saved as separate CSV files. Phase 5, including the paired bootstrap interval, has code but no executed output in the notebook. For this reason, the aggregate Phase 4 values are described as provisional throughout the draft.", bold_lead="Evidence status.")
    doc.add_page_break()


def build_paper(doc, fig1, fig2):
    # Abstract
    doc.add_heading("Abstract", level=1)
    abstract = (
        "Data-driven photovoltaic (PV) power models often achieve low error when their training and test sets contain observations from the same installation, but this evaluation does not answer a practical cold-start question: can one model predict the output of a PV system whose historical power data were never used in training? This study investigates whether physically derived, system-independent features improve such cross-system generalization. The analysis uses the open HKUST rooftop PV dataset, which contains three years of power and meteorological measurements from 60 campus installations. After restricting the analysis to 37 SolarEdge systems, localizing timestamps, filtering nighttime observations, constructing physics features, and applying quality-control rules, 1,370,191 modeling-ready observations remained. Four predictors were compared under five-fold grouped cross-validation in which entire PV systems, rather than individual timestamps, were held out: a training-mean baseline, a calibrated global-horizontal-irradiance baseline, weather-only XGBoost, and an otherwise identical XGBoost model augmented with solar zenith, solar azimuth, clear-sky irradiance, clear-sky index, and estimated cell temperature. The executed notebook reports macro normalized root-mean-square errors of 0.2160, 0.1237, 0.0919, and 0.0790, respectively. Thus, the physics-enhanced model reduced macro normalized error by 14.0% and mean absolute error by 15.0% relative to the weather-only model. However, the hybrid model also showed a larger cross-system standard deviation and a slightly worse maximum system error. The results provide preliminary evidence that compact physics-derived features improve average transfer to unseen systems within one campus, while also showing that aggregate gains do not guarantee uniform improvement. A paired confidence interval and cross-climate validation remain necessary before the conclusion can be treated as final."
    )
    add_body(doc, abstract)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("Keywords: ")
    set_run_font(r, bold=True)
    r2 = p.add_run("photovoltaics; physics-informed machine learning; XGBoost; domain generalization; cold-start prediction; rooftop solar")
    set_run_font(r2, italic=True)

    # 1 Introduction
    doc.add_heading("1. Introduction", level=1)
    add_body(doc, "Rooftop photovoltaic generation is inherently variable because the amount of electricity produced at a given moment depends on solar geometry, atmospheric attenuation, clouds, temperature, wind, installation capacity, orientation, and equipment condition. This variability complicates monitoring and planning for distributed solar fleets. Prediction models can help estimate expected production, detect abnormal behavior, and support grid and building operations, but their usefulness depends on how they are evaluated. A model that performs well on randomly held-out timestamps from a familiar system may be learning that system's characteristic scale and response rather than a relationship that transfers to a new installation.")
    add_body(doc, "The cold-start setting is especially important for new or sparsely monitored rooftop systems. A newly commissioned system has no long history of measured output from which to learn site-specific behavior. In such a setting, a useful model must transfer relationships learned from other installations. Prior research has shown that a single XGBoost model can nowcast individual yield across a large residential fleet and that models trained on geographically aggregated data can generalize to individual systems (Costa, 2022; Grzebyk et al., 2023). More recent work has addressed zero-sample PV prediction using explicit domain-generalization architectures (Liu et al., 2025). These studies establish that cross-system prediction is both possible and technically distinct from ordinary within-system forecasting.")
    add_body(doc, "Physical knowledge offers one plausible route to improved transfer. Measured global horizontal irradiance (GHI) captures current sunlight, but the same GHI value can occur under different solar positions and clear-sky expectations. Solar zenith and azimuth encode geometry; clear-sky GHI provides a time- and location-specific atmospheric reference; the clear-sky index expresses observed irradiance relative to that reference; and an estimated cell temperature represents the thermal conditions that influence PV conversion efficiency. These variables are derived from relationships that apply across systems and therefore may help a model rely less on installation-specific statistical patterns.")
    add_body(doc, "This paper evaluates that hypothesis using a controlled feature-ablation experiment. Two XGBoost regressors use the same target, grouped folds, algorithm, hyperparameters, and random seed. Model A receives four measured weather variables. Model B receives the same four variables plus five physics-derived features. Because the only planned treatment difference is the feature set, the error difference between the two models estimates the predictive value of those physics features under the chosen evaluation protocol. XGBoost is appropriate for this comparison because tree boosting captures nonlinear interactions in tabular data and has demonstrated strong performance in large PV fleets (Chen & Guestrin, 2016; Grzebyk et al., 2023).")
    add_body(doc, "The study makes three contributions. First, it implements a leakage-resistant evaluation in which all observations from a test system are excluded from training. Second, it compares physics-only, data-driven, and hybrid approaches on identical held-out systems. Third, it reports system-level macro metrics so that installations with more observations do not dominate the headline result. The intended contribution is deliberately narrower than proposing a new neural architecture: it asks whether a small, interpretable physics feature set adds measurable value to a standard tabular learner.")
    add_body(doc, "The executed notebook reports a 14.0% reduction in macro normalized root-mean-square error (nRMSE) for the hybrid model relative to the weather-only model. That result is promising but not yet sufficient on its own. The system-level Phase 4 results must be saved and the paired uncertainty analysis must be executed before final submission. Moreover, all systems share one campus, climate, and weather station, so the experiment evaluates cross-system generalization rather than cross-climate generalization. The paper therefore treats the current result as evidence within a bounded empirical setting, not as proof that the model will transfer to arbitrary locations.")

    # 2 Problem definition
    doc.add_heading("2. Problem Definition and Hypotheses", level=1)
    add_body(doc, "The prediction unit is one daylight observation for one rooftop PV system. For system s at timestamp t, the response variable is the measured AC power P divided by the system's installed capacity C. Capacity normalization places systems of different sizes on a comparable scale:")
    add_equation(doc, "y(s,t) = P(s,t) / C(s)")
    add_body(doc, "The normalized target is approximately interpretable as a fraction of installed capacity. It is not a capacity factor averaged over a long period; it is an instantaneous normalized power value. Rows outside the accepted range from -0.001 to 1.2 are marked impossible and excluded from modeling. The small negative tolerance accommodates numerical or measurement noise, while the upper bound allows modest output above nominal capacity without admitting extreme values.")
    add_body(doc, "The central evaluation condition is zero historical power data from a test system. Let S_train be the systems assigned to a training fold and S_test the systems assigned to its test fold, with S_train intersect S_test equal to the empty set. A regressor is fit using all modeling-ready rows from S_train and evaluated on all modeling-ready rows from S_test. This differs from random row-level splitting, which would place different timestamps from the same system in both sets.")
    add_body(doc, "Two feature sets define the controlled comparison. Model A uses measured GHI, ambient temperature, wind speed, and relative humidity. Model B uses those same variables plus apparent solar zenith, solar azimuth, clear-sky GHI, clear-sky index, and a GHI-based Faiman cell-temperature estimate. The null and alternative hypotheses are:")
    add_bullet(doc, "H0: Adding the five physics-derived features does not reduce mean system-level nRMSE relative to the weather-only XGBoost model.")
    add_bullet(doc, "H1: Adding the five physics-derived features reduces mean system-level nRMSE relative to the weather-only XGBoost model.")
    add_body(doc, "The primary outcome is system-level nRMSE averaged equally across held-out systems. Secondary outcomes are normalized mean absolute error (MAE) and mean bias error (MBE). The primary comparison is paired because Models A and B are evaluated on the same systems and folds. The analysis is predictive, not causal: it evaluates whether adding a feature set improves out-of-sample error within this pipeline. It also predicts power conditional on contemporaneous measured weather; it does not forecast future weather or evaluate an operational day-ahead forecast.")

    # 3 Related research
    doc.add_heading("3. Related Research", level=1)
    doc.add_heading("3.1 Data-driven PV prediction", level=2)
    add_body(doc, "Machine-learning models are attractive for PV prediction because the mapping from meteorological inputs to power is nonlinear and installation-dependent. XGBoost constructs an additive ensemble of decision trees and includes regularization and subsampling mechanisms that support scalable learning from large tabular datasets (Chen & Guestrin, 2016). In a study of 1,102 residential systems, Grzebyk et al. (2023) selected XGBoost for individual yield nowcasting because of its performance and practical simplicity. Their results demonstrated that one fleet-level model can outperform commercial software for individual systems, supporting the use of a shared learner rather than a separate model per installation.")
    add_body(doc, "Deep sequence models provide another path. Costa (2022) compared LSTM, convolutional, and hybrid convolutional-LSTM networks for household PV generation and evaluated whether models trained on regionally aggregated data could predict individual systems. The reported results supported regional-to-individual transfer. That study is relevant because it treats generalization as an explicit evaluation question, rather than assuming that accuracy on randomly held-out observations implies transfer.")

    doc.add_heading("3.2 Physics-based and physics-derived representations", level=2)
    add_body(doc, "A physics-based PV workflow typically transforms time and location into solar position, estimates clear-sky or actual irradiance components, estimates module temperature, and then maps irradiance and temperature to electrical output. The open-source pvlib library provides modular implementations of solar-position, irradiance, thermal, and electrical models, allowing researchers to use physical calculations either as complete models or as components of hybrid pipelines (Anderson et al., 2023). This study uses pvlib to construct physically meaningful inputs while leaving the final weather-to-power mapping to XGBoost.")
    add_body(doc, "The clear-sky reference uses the Ineichen-Perez formulation. Ineichen and Perez (2002) developed an airmass-independent formulation of the Linke turbidity coefficient and associated clear-sky irradiance models. Dividing observed GHI by modeled clear-sky GHI produces a clear-sky index that partially separates predictable solar geometry from cloud-related attenuation. This representation can make the meaning of a given irradiance value more comparable across hours and seasons.")
    add_body(doc, "Temperature is another important physical pathway because PV conversion efficiency changes with cell temperature. The Faiman model estimates module or cell temperature using incident irradiance, ambient temperature, wind speed, and empirical heat-loss coefficients (Faiman, 2008). The current implementation applies the default coefficients u0 = 25.0 and u1 = 6.84. However, it substitutes GHI for plane-of-array irradiance because the system-level records may aggregate arrays with multiple or uncertain orientations. This approximation creates a useful thermal proxy but should not be interpreted as a fully specified module temperature model.")

    doc.add_heading("3.3 Generalization to unseen PV systems", level=2)
    add_body(doc, "The literature increasingly distinguishes interpolation within familiar sites from prediction on unknown systems. Liu et al. (2025) proposed a generative adversarial domain-generalization network for zero-sample prediction and evaluated nine geographically and capacity-diverse PV systems. Their results show that unseen-system prediction is now an active research area rather than an entirely untested problem. The present study therefore does not claim to be the first cross-system PV model. Its contribution is a transparent controlled ablation on a larger set of installations within one campus, designed to quantify the incremental value of a compact physics feature set in a standard XGBoost pipeline.")
    add_body(doc, "This distinction matters for novelty. Complex domain-generalization models may maximize predictive accuracy, but they combine architecture, augmentation, and representation changes. A simpler paired comparison can answer a different question: how much improvement is obtained by adding five physically interpretable variables while holding the learner fixed? A positive answer would support physics-derived feature engineering as a low-complexity baseline for future domain-generalization research. A negative answer would be equally informative because it would show that physical transformations do not automatically improve transfer.")

    doc.add_heading("3.4 Research gap", level=2)
    add_body(doc, "Existing studies demonstrate fleet-level prediction, regional generalization, and advanced zero-sample domain generalization, but they do not determine the value of this exact physics feature set under a system-grouped ablation on the HKUST rooftop dataset. The research gap is therefore dataset- and design-specific: the incremental cross-system benefit of solar position, clear-sky context, and a thermal proxy has not been quantified for these campus rooftops under an otherwise identical XGBoost configuration. The study addresses that bounded gap and avoids a broader priority claim that the literature cannot support.")

    # 4 Dataset
    doc.add_heading("4. Dataset and Data Preparation", level=1)
    doc.add_heading("4.1 Source dataset and analytical subset", level=2)
    add_body(doc, "Lin et al. (2025) released a three-year rooftop PV dataset collected at the Hong Kong University of Science and Technology. The source contains inverter-level generation measurements from 60 grid-connected rooftop stations at five-minute resolution, on-site meteorological measurements at one-minute resolution, and system metadata. The campus is located in Hong Kong's humid subtropical climate. The source paper explicitly notes that the single-location climate may limit the generalizability of models trained on the data.")
    add_body(doc, "The current project begins from a previously combined site-level table with 2,756,516 rows. Its measurements have been aligned to 15-minute timestamps. The analysis retains 37 SolarEdge systems for which the required power, capacity, and weather fields are available in the combined table. Although the source dataset contains 60 stations, this paper's claims apply only to the 37-system analytical subset. The reasons that the remaining systems are out of scope should be documented in the final version, particularly if the SMA systems are analyzed separately.")
    add_table(
        doc,
        ["Quantity", "Value", "Interpretation"],
        [
            ("Source period", "2021-2023", "Three years of campus operation"),
            ("Source stations", "60", "All rooftop stations described by Lin et al. (2025)"),
            ("Analyzed systems", "37", "SolarEdge systems in the combined modeling table"),
            ("Raw combined rows", "2,756,516", "Before daylight filtering"),
            ("Daylight rows", "1,387,682", "Apparent zenith below 90 degrees"),
            ("Modeling-ready rows", "1,370,191", "After exclusion-quality flags"),
            ("Flagged daylight rows", "17,491", "Retained for audit but excluded from modeling"),
        ],
        [2700, 1700, 4960],
        numeric_cols=(1,),
        caption="Table 1. Dataset scope after preprocessing",
        source="Source: Executed outputs in code.ipynb; source-dataset description from Lin et al. (2025).",
    )

    doc.add_heading("4.2 Timestamp handling and daylight filtering", level=2)
    add_body(doc, "Timestamps are localized to the Asia/Hong_Kong time zone rather than converted from another zone. Solar position is calculated for the campus coordinates 22.3364 degrees north, 114.2654 degrees east, at an assumed altitude of 60 m. Apparent solar zenith and azimuth are merged back to all system rows by timestamp. Only observations with apparent zenith below 90 degrees are retained. Removing nighttime rows avoids a misleading reduction in error from long sequences in which both observed and predicted power are trivially zero.")
    add_body(doc, "The daylight filter retained 1,387,682 of 2,756,516 rows. The maximum retained zenith was 89.9997 degrees. Because the same campus weather station is shared across systems, the code also verifies that weather columns do not vary among stations at a common timestamp. This many-system/one-weather-source structure is useful for isolating system differences but limits claims about geographic transfer.")

    doc.add_heading("4.3 Physics feature construction", level=2)
    add_body(doc, "Clear-sky GHI is computed for every retained timestamp with pvlib's Ineichen model. The clear-sky index is the measured GHI divided by the corresponding clear-sky GHI and is capped at 2.0 by the library function. A GHI-based cell-temperature proxy is computed with the Faiman model using ambient temperature, wind speed, u0 = 25.0, and u1 = 6.84. Finally, normalized power is calculated as measured power divided by the summed installed capacity for the system. The engineered values are merged by timestamp or computed row by row, without using target-system power information to construct the weather or solar-geometry inputs.")

    doc.add_heading("4.4 Quality control", level=2)
    add_body(doc, "The preprocessing pipeline is non-destructive: quality-control flags remain in the cleaned table, while the modeling-ready indicator determines which rows enter training and evaluation. A daytime fault zero is defined as zero power when GHI is at least 200 W/m2. A stuck measurement is an exactly repeated positive power or weather value lasting at least 16 consecutive 15-minute samples, equivalent to four hours. Additional rules flag missing model variables, nonpositive capacity, negative GHI or wind speed, relative humidity outside 0-100%, and normalized power outside -0.001 to 1.2.")
    add_body(doc, "The executed quality summary identified 16,991 rows with missing values and 500 daytime fault-zero rows. It detected no impossible values and no four-hour stuck power or weather runs. Seven hundred thirty-one rows occurred after a detected data gap, but this informational flag did not exclude rows because the tabular models do not use lagged features. In total, 17,491 daylight rows were excluded from modeling. The exact thresholds are transparent and reproducible, but they remain analyst choices; sensitivity to the fault-zero and stuck-sensor thresholds is a potential robustness check.")

    doc.add_paragraph().add_run()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(fig1), width=Inches(6.35))
    add_caption(doc, "Figure 1. Experimental design. All timestamps from each held-out PV system remain outside the training set. Models A and B use identical folds and XGBoost settings, so their paired error difference isolates the added physics-derived features. Source: Project notebook design.")

    # 5 Methods
    doc.add_heading("5. Methodology", level=1)
    doc.add_heading("5.1 Feature-ablation design", level=2)
    add_body(doc, "The principal experiment is a controlled feature ablation. Model A and Model B use the same regression algorithm, objective, hyperparameters, capacity-normalized target, model-ready rows, grouped folds, and random seed. Model A contains four measured weather features. Model B adds five deterministic physics-derived features. Plane-of-array irradiance is deliberately excluded because a single system-level record may combine arrays with different tilts or azimuths, making one orientation estimate potentially misleading.")
    add_table(
        doc,
        ["Feature", "Model A", "Model B", "Rationale"],
        [
            ("Measured GHI", "Yes", "Yes", "Current solar resource"),
            ("Ambient temperature", "Yes", "Yes", "Environmental and thermal context"),
            ("Wind speed", "Yes", "Yes", "Convective cooling context"),
            ("Relative humidity", "Yes", "Yes", "Atmospheric context"),
            ("Apparent zenith", "No", "Yes", "Solar elevation geometry"),
            ("Solar azimuth", "No", "Yes", "Solar direction geometry"),
            ("Clear-sky GHI", "No", "Yes", "Time/location-specific irradiance reference"),
            ("Clear-sky index", "No", "Yes", "Observed irradiance relative to clear sky"),
            ("GHI-based cell temperature", "No", "Yes", "Approximate thermal state"),
        ],
        [3000, 1050, 1050, 4260],
        numeric_cols=(1, 2),
        caption="Table 2. Controlled feature sets",
        source="Source: MODEL_A_FEATURES and MODEL_B_FEATURES in code.ipynb.",
    )

    doc.add_heading("5.2 Reference baselines", level=2)
    add_body(doc, "The training-mean baseline predicts the mean normalized power of the training rows in each fold. It is intentionally weak but establishes the error of a constant predictor under the same split. The GHI physics baseline first computes a solar signal equal to clear-sky GHI divided by 1,000 and multiplied by the clear-sky index. Because the index reconstructs observed GHI up to the clipping rule, the signal is a capacity-normalized irradiance proxy. A nonnegative scalar is fit through the origin using only training rows, and predictions are clipped between 0 and 1.2. The fitted scalar ranges from 0.6254 to 0.6382 across the five folds.")
    add_body(doc, "This baseline is not a full PVWatts or system-specific physical model. It does not use system tilt, azimuth, module type, inverter efficiency, or loss parameters. Its purpose is to test how far a simple physical irradiance relationship can go without target-system calibration.")

    doc.add_heading("5.3 XGBoost specification", level=2)
    add_body(doc, "Both machine-learning models use XGBoost regression with squared-error loss. Gradient-boosted trees sequentially add trees that correct residual error while regularization and subsampling reduce overfitting (Chen & Guestrin, 2016). Tree models do not require feature standardization, so scaling is disabled. No separate hyperparameter tuning is performed for Models A and B; tuning them independently would confound the feature comparison.")
    add_table(
        doc,
        ["Parameter", "Value", "Purpose"],
        [
            ("objective", "reg:squarederror", "Squared-error regression"),
            ("n_estimators", "1,000", "Number of boosting rounds"),
            ("learning_rate", "0.05", "Shrinkage per tree"),
            ("max_depth", "6", "Maximum interaction depth"),
            ("min_child_weight", "10", "Minimum child-node weight"),
            ("subsample", "0.8", "Row subsampling"),
            ("colsample_bytree", "0.8", "Feature subsampling per tree"),
            ("reg_alpha", "0.0", "L1 regularization"),
            ("reg_lambda", "1.0", "L2 regularization"),
            ("tree_method", "hist", "Histogram-based training"),
            ("random_state", "42", "Deterministic split/model seed"),
        ],
        [2600, 2200, 4560],
        numeric_cols=(1,),
        caption="Table 3. Shared XGBoost configuration",
        source="Source: PHASE4_MODEL_CONFIG in code.ipynb. Exact package versions remain [TO COMPLETE].",
    )

    doc.add_heading("5.4 Leave-systems-out evaluation", level=2)
    add_body(doc, "Five-fold GroupKFold cross-validation is used with system ID as the group variable. The random seed is 42 and shuffling is enabled. Each fold contains approximately 29-30 training systems and 7-8 test systems. Every system appears in a test fold exactly once. Explicit assertions verify that the training-system and test-system sets are disjoint and that each expected system receives one evaluation result. This design follows the grouped-validation principle implemented in scikit-learn (Pedregosa et al., 2011).")
    add_body(doc, "Predictions are evaluated at two levels. Fold-level metrics summarize all rows in a fold, while system-level metrics are calculated separately for each held-out installation. Headline values are macro means across the 37 system results, giving each system equal weight regardless of its number of observations. This choice aligns the estimand with the question: expected performance on an unseen system, not expected error on a randomly selected timestamp from the pooled dataset.")

    doc.add_heading("5.5 Performance metrics", level=2)
    add_body(doc, "Because the target is already divided by installed capacity, all errors are expressed as fractions of capacity. For observed normalized power y_i and prediction y-hat_i over N observations for a system, the metrics are:")
    add_equation(doc, "nRMSE = √[(1/N) Σᵢ (ŷᵢ - yᵢ)²]")
    add_equation(doc, "MAE = (1/N) Σᵢ |ŷᵢ - yᵢ|")
    add_equation(doc, "MBE = (1/N) Σᵢ (ŷᵢ - yᵢ)")
    add_body(doc, "Lower nRMSE and MAE indicate greater accuracy. MBE is signed: positive values indicate average overprediction and negative values indicate average underprediction. The relative error reduction of Model B against Model A is 100 times the difference between Model A and Model B error divided by Model A error.")

    doc.add_heading("5.6 Leakage controls and reproducibility", level=2)
    add_body(doc, "The code contains several safeguards. All rows from a system remain in one fold. Preprocessing that can learn from data is fitted on training rows only. Models A and B use identical group assignments. The cleaned table retains time-zone information and verifies unique system-timestamp rows. Model-ready rows contain no missing model variables or active exclusion flags. Fixed random seeds support deterministic reruns.")
    add_body(doc, "Reproducibility is not yet complete at the repository level. The notebook's active input path points to /content/site_level_dataset_modified_combined.csv, and the source combined table and cleaned parquet file are not present in the current project folder. The requirements file lists dependencies without version pins. Phase 3 results are saved, but Phase 4 and Phase 5 result tables are absent even though Phase 4 output is embedded in the notebook. Before submission, the project should document dataset retrieval, use a portable input path, pin software versions, and commit or archive all non-sensitive derived result tables.")

    doc.add_heading("5.7 Planned statistical comparison", level=2)
    add_body(doc, "The appropriate uncertainty analysis is paired at the system level. For each held-out system, the nRMSE difference is computed as Model B minus Model A. A deterministic paired bootstrap can resample the 37 systems with replacement, calculate the macro difference and relative reduction for each resample, and use the 2.5th and 97.5th percentiles as a 95% interval. Resampling rows would understate uncertainty because adjacent timestamps from the same system are dependent. The code for this analysis exists in Phase 5, but the cell has not been executed in the available notebook. The final paper must report the interval and the number of systems on which Model B improves: [TO COMPLETE].")

    # 6 Results
    doc.add_heading("6. Experimental Results", level=1)
    doc.add_heading("6.1 Physics baseline substantially improves on a constant predictor", level=2)
    add_body(doc, "The calibrated GHI baseline reduced macro nRMSE from 0.2160 to 0.1237, a relative reduction of 42.7%. It reduced macro MAE from 0.1811 to 0.0776, or 57.2%. The GHI baseline had lower nRMSE than the training-mean predictor on all 37 held-out systems. This result is a useful validation of the experimental setup: a simple irradiance-based relationship contains transferable information that the constant baseline lacks.")
    add_body(doc, "The physics baseline nevertheless underpredicted on average, with macro MBE of -0.0170. It also remained materially less accurate than either XGBoost model. Therefore, physics alone provides a meaningful reference but does not capture all nonlinear relationships between weather and PV output in the analyzed systems.")

    doc.add_heading("6.2 Physics features reduce average XGBoost error by approximately 14%", level=2)
    add_body(doc, "The executed Phase 4 notebook output reports macro nRMSE of 0.0919 for the weather-only model and 0.0790 for the physics-enhanced model. The absolute difference is -0.0129 of installed capacity, corresponding to a 14.0% relative reduction. Macro MAE falls from 0.0614 to 0.0522, a 15.0% relative reduction. Mean bias also moves closer to zero, from 0.0010 for Model A to -0.0004 for Model B.")
    add_table(
        doc,
        ["Model", "Macro nRMSE", "SD", "Macro MAE", "Macro MBE"],
        [
            ("Training mean", "0.2160", "0.0185", "0.1811", "-0.0013"),
            ("Calibrated GHI physics", "0.1237", "0.0146", "0.0776", "-0.0170"),
            ("Model A: weather-only XGBoost", "0.0919", "0.0190", "0.0614", "0.0010"),
            ("Model B: physics-enhanced XGBoost", "0.0790", "0.0228", "0.0522", "-0.0004"),
        ],
        [3300, 1600, 1300, 1600, 1560],
        numeric_cols=(1, 2, 3, 4),
        caption="Table 4. Macro performance across 37 held-out PV systems",
        source="Source: Phase 3 CSV files and executed Phase 4 output in code.ipynb. Phase 4 values are provisional until the result CSVs are saved and rechecked.",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(fig2), width=Inches(6.35))
    add_caption(doc, "Figure 2. Macro nRMSE across four models. Errors are calculated on capacity-normalized power for each held-out system and then averaged equally across 37 systems. Lower values indicate better cross-system prediction. The hybrid value is 14.0% below the weather-only value. Source: Project Phase 3 results and executed Phase 4 notebook output.")

    doc.add_heading("6.3 The average gain may not be uniform across systems", level=2)
    add_body(doc, "Aggregate improvement is not the only relevant result. Model B's system-level nRMSE standard deviation is 0.0228, compared with 0.0190 for Model A. Its median is lower, 0.0720 versus 0.0879, suggesting that a typical system benefits. However, the maximum system nRMSE is slightly worse for Model B, 0.1676 versus 0.1598. These values imply heterogeneous effects: physics features improve the macro average and median but may harm at least one difficult system.")
    add_body(doc, "This heterogeneity is scientifically important because the practical question concerns reliability on a new system, not only average accuracy. The final analysis should report the paired difference for every system, the number of systems on which Model B wins, and a scatterplot against the identity line. Those values cannot be reconstructed reliably from the aggregate notebook output and remain [TO COMPLETE].")

    doc.add_heading("6.4 Statistical uncertainty remains pending", level=2)
    add_body(doc, "The observed 14.0% reduction is a point estimate over 37 systems. Without the paired bootstrap interval, the paper cannot determine how precisely that reduction estimates expected performance across similar campus systems. The planned 95% confidence interval is [TO COMPLETE: lower bound, upper bound]. Until that analysis is executed, the evidence supports the descriptive statement that Model B performed better on average in these folds, but not a final inferential claim about the broader population of PV systems.")

    doc.add_heading("6.5 Robustness and interpretation analyses", level=2)
    add_body(doc, "Two planned analyses are not yet complete. First, the notebook includes a gated hook for evaluating SMA systems as a robustness subset, but the cleaned SMA input has not been prepared. Second, no feature-importance analysis is available to determine whether the hybrid model relies primarily on clear-sky index, solar geometry, or the temperature proxy. These analyses are not required to establish the main feature-ablation result, but they would strengthen the explanation of why Model B improves and whether the effect extends beyond the SolarEdge subset.")

    # 7 Discussion
    doc.add_heading("7. Discussion", level=1)
    doc.add_heading("7.1 Why physics-derived features may improve transfer", level=2)
    add_body(doc, "The average improvement is consistent with the idea that physical transformations provide a more stable representation across systems and seasons. Raw GHI indicates the instantaneous resource, but it does not directly state whether that irradiance is high or low relative to what is possible at the current solar position. Clear-sky GHI and clear-sky index add that context. Solar zenith and azimuth encode deterministic daily and seasonal geometry. The cell-temperature proxy combines irradiance, ambient temperature, and wind into a variable more directly related to conversion efficiency than any one weather measurement alone.")
    add_body(doc, "Tree ensembles can in principle learn interactions among time, irradiance, temperature, and wind if those variables are supplied. However, Model A does not include time-of-day or day-of-year features. The physics features therefore do more than add domain knowledge: solar position and clear-sky irradiance also encode time and season in a physically meaningful way. The current experiment estimates the value of the complete five-feature bundle, not the isolated value of each physical mechanism. A follow-up ablation should add solar geometry, clear-sky variables, and temperature sequentially to separate these contributions.")

    doc.add_heading("7.2 Comparison with related work", level=2)
    add_body(doc, "The results align with prior evidence that shared models can predict individual PV systems (Costa, 2022; Grzebyk et al., 2023). They also complement the domain-generalization approach of Liu et al. (2025). Whereas that work proposes a complex adversarial framework across nine diverse systems, this study uses a standard learner and 37 systems from one campus to isolate feature representation. The approaches are not direct competitors because their datasets, inputs, prediction horizons, and metrics differ. The contribution here is interpretability and experimental control rather than architectural novelty.")
    add_body(doc, "The comparison with the GHI baseline also clarifies the role of hybrid modeling. A simple irradiance-to-power relationship transfers better than a constant predictor, but both XGBoost models improve substantially on it. This suggests that the data-driven learner captures nonlinearities and operational differences not represented by the single physics scale. The hybrid model then improves further, indicating that statistical flexibility and physical context are complementary in this setting.")

    doc.add_heading("7.3 Practical meaning of the reported error", level=2)
    add_body(doc, "An nRMSE of 0.0790 means that the root-mean-square prediction error is approximately 7.9% of installed capacity for the average held-out system. For a 100 kW system, this normalized value corresponds to an RMSE of roughly 7.9 kW under the same daylight and contemporaneous-weather conditions, although system-specific error may be higher or lower. The reduction from 0.0919 to 0.0790 corresponds to approximately 1.29 percentage points of installed capacity.")
    add_body(doc, "That interpretation should not be confused with day-ahead operational forecasting. The models receive measured contemporaneous irradiance and weather, so they answer what output is expected given current conditions. If deployed with weather forecasts, additional error from the weather forecast would enter the system. The current model may therefore be most directly useful for nowcasting, performance benchmarking, and fault-screening applications.")

    doc.add_heading("7.4 Why some systems may not improve", level=2)
    add_body(doc, "The slightly worse maximum error for Model B suggests that physical proxies can be imperfect. The same clear-sky and solar-position variables are applied to every system, but rooftops differ in orientation, shading, module technology, inverter behavior, commissioning dates, and maintenance conditions. The GHI-based temperature estimate substitutes horizontal irradiance for plane-of-array irradiance, which may be less representative for steep or multi-orientation arrays. A tree model can also overuse a derived feature that correlates well for most systems but not for an atypical one.")
    add_body(doc, "System-level diagnostics should therefore compare improvement against available metadata such as capacity, optimizer type, orientation complexity, and missing-data rate. If the hybrid model performs worse mainly on multi-orientation or shaded systems, the result would suggest where more detailed physical metadata is necessary. If losses are unrelated to known metadata, the remaining variation may reflect unobserved equipment or data-quality differences.")

    # 8 Limitations
    doc.add_heading("8. Limitations, Uncertainty, and Threats to Validity", level=1)
    add_body(doc, "Several limitations bound the conclusions. First, all analyzed systems share one campus, one climate, and one weather station. Holding out systems tests variation in equipment and rooftop configuration, but it does not test geographic, meteorological, or sensor-domain shift. The experiment supports cross-system generalization within the HKUST environment, not cross-climate generalization.")
    add_body(doc, "Second, the analysis is conditional power prediction rather than a complete forecast. Contemporaneous measured weather is supplied to the models. Future work using numerical weather predictions would need to propagate weather forecast error and define an explicit horizon. Terms such as nowcasting or conditional prediction are more accurate than day-ahead forecasting for the current design.")
    add_body(doc, "Third, physical feature accuracy is constrained by metadata. Plane-of-array irradiance is omitted because the site-level power series can aggregate arrays with uncertain or mixed orientations. The Faiman temperature proxy therefore uses GHI, and its default coefficients were originally derived for open-rack modules under different conditions (Faiman, 2008). These choices favor robustness and availability over system-specific physical fidelity.")
    add_body(doc, "Fourth, the 15-minute rows within a system are temporally dependent. Reporting millions of rows does not imply millions of independent observations. System-level macro evaluation and paired system resampling partially address this issue, but the effective sample size for generalization is closer to 37 systems than to 1.37 million timestamps. Claims of statistical certainty must reflect that grain.")
    add_body(doc, "Fifth, the current result is not fully reproducible from the checked project folder. The raw combined input and cleaned parquet are not present, dependency versions are unpinned, and Phase 4/5 CSV outputs are missing. The embedded notebook output provides evidence for drafting but should not be the sole record for a final paper. A clean rerun from documented inputs is required.")
    add_body(doc, "Finally, the model configuration was selected before the paired comparison but not established through a nested tuning procedure. Holding the configuration fixed is appropriate for isolating feature value, yet it does not show that either feature set is optimally tuned. The result should be interpreted as the improvement under one reasonable common XGBoost configuration.")

    # 9 Future work
    doc.add_heading("9. Future Work", level=1)
    add_body(doc, "The first priority is to complete the existing analysis rather than add model complexity. Phase 5 should save the paired system results, calculate the bootstrap interval, report system win counts, and generate the paired scatterplot. The preprocessing and modeling pipeline should then be rerun from a portable dataset path with pinned versions and fixed seeds.")
    add_body(doc, "A second priority is robustness across hardware and location. The planned SMA subset can test whether the conclusion changes for systems without panel-level optimizers. A stronger external test would train on Hong Kong and evaluate zero-shot performance on a public dataset from a different climate, such as Alice Springs. Such a test would move the claim from cross-system to cross-climate generalization.")
    add_body(doc, "A third direction is mechanism-focused ablation. Adding physics features in groups would distinguish the contribution of geometry, clear-sky normalization, and temperature. Permutation importance or SHAP analysis could show which features the hybrid uses, but those tools should be interpreted as predictive explanations rather than causal effects.")
    add_body(doc, "Finally, physics-as-features can be compared with physics-guided residual correction. In the latter design, a physical model produces an initial prediction and machine learning models its residual error. Comparing these two hybridization strategies under the same held-out systems would test whether physical structure is more transferable as an input representation or as a baseline model.")

    # 10 Conclusion
    doc.add_heading("10. Overall Conclusion", level=1)
    add_body(doc, "This study asks whether a compact set of physics-derived variables improves PV power prediction for rooftop systems that are entirely absent from training. Under five-fold leave-systems-out evaluation on 37 HKUST SolarEdge systems, the executed notebook reports that weather-only XGBoost achieves macro nRMSE of 0.0919, while the otherwise identical physics-enhanced model achieves 0.0790. The difference corresponds to a 14.0% relative reduction in normalized error. The hybrid also reduces MAE by 15.0% and brings mean bias close to zero.")
    add_body(doc, "These results support the practical value of solar position, clear-sky context, and an approximate cell-temperature feature for average cross-system transfer within one campus. At the same time, the hybrid's larger error spread and slightly worse maximum error warn that the gain may not apply uniformly. Because the paired confidence interval and system win count have not yet been generated, the final inferential conclusion remains provisional.")
    add_body(doc, "The most defensible current conclusion is therefore bounded: physics-derived features improved average held-out-system prediction in this specific dataset and controlled XGBoost pipeline. Completing the paired analysis and testing an external climate are necessary before claiming general transfer to unseen rooftop PV systems.")

    # Availability and AI disclosure
    doc.add_heading("Data and Code Availability", level=1)
    add_body(doc, "The source dataset is described by Lin et al. (2025) and is publicly available through the repository linked in that publication. Analysis code is contained in the project notebook. Final repository URL, release tag, license, dataset retrieval instructions, and software environment: [TO COMPLETE]. The cleaned data should be distributed only if permitted by the source dataset's license; otherwise, provide a reproducible download-and-build script.")

    doc.add_heading("AI Assistance Disclosure", level=1)
    add_body(doc, "OpenAI Codex was used on July 21, 2026 to create an initial paper structure, draft prose, format the Word document, and identify places where evidence was incomplete. The tool was given the course syllabus, the local project notebook, saved result tables, and primary literature metadata. It did not independently run the unavailable raw-data pipeline or verify the unexecuted Phase 5 analysis. The student author is responsible for checking every claim, revising the text, documenting the interaction in the manner required by the Pioneer Research Program, and ensuring that the final paper represents the student's own reasoning and research. Suggested citation: OpenAI. (2026). Codex [Large language model]. https://openai.com/codex/")

    # References
    doc.add_heading("References", level=1)
    refs = [
        "Anderson, K. S., Hansen, C. W., Holmgren, W. F., Jensen, A. R., Mikofski, M. A., & Driesse, A. (2023). pvlib python: 2023 project update. Journal of Open Source Software, 8(92), 5994. https://doi.org/10.21105/joss.05994",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785-794). Association for Computing Machinery. https://doi.org/10.1145/2939672.2939785",
        "Costa, R. L. de C. (2022). Convolutional-LSTM networks and generalization in forecasting of household photovoltaic generation. Engineering Applications of Artificial Intelligence, 116, 105458. https://doi.org/10.1016/j.engappai.2022.105458",
        "Faiman, D. (2008). Assessing the outdoor operating temperature of photovoltaic modules. Progress in Photovoltaics: Research and Applications, 16(4), 307-315. https://doi.org/10.1002/pip.813",
        "Grzebyk, D., Alcañiz, A., Donker, J., Zeman, M., Ziar, H., & Isabella, O. (2023). Individual yield nowcasting for residential PV systems. Solar Energy, 251, 325-336. https://doi.org/10.1016/j.solener.2023.01.036",
        "Ineichen, P., & Perez, R. (2002). A new airmass independent formulation for the Linke turbidity coefficient. Solar Energy, 73(3), 151-157. https://doi.org/10.1016/S0038-092X(02)00045-2",
        "Lin, Z., Zhou, Q., Wang, Z., Wang, C., Bookhart, D. B., & Leung-Shea, M. (2025). A high-resolution three-year dataset supporting rooftop photovoltaics (PV) generation analytics. Scientific Data, 12, Article 63. https://doi.org/10.1038/s41597-025-04397-y",
        "Liu, S., Qi, Y., Li, D., Liu, L., Wang, S., Fernandez, C., & Gao, X. (2025). Adversarial multi-source domain generalization approach for power prediction in unknown photovoltaic systems. Applied Soft Computing, 181, 113495. https://doi.org/10.1016/j.asoc.2025.113495",
        "OpenAI. (2026). Codex [Large language model]. https://openai.com/codex/",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830. https://www.jmlr.org/papers/v12/pedregosa11a.html",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        set_run_font(r, size=10.5)

    # Appendix
    doc.add_page_break()
    doc.add_heading("Appendix A. Final-analysis completion table", level=1)
    add_body(doc, "This appendix is part of the working draft and should be removed or converted into a methods supplement before submission.")
    add_table(
        doc,
        ["Required item", "Current status", "Final-paper action"],
        [
            ("Phase 3 baselines", "Complete and saved", "Reconfirm after clean rerun"),
            ("Phase 4 model summaries", "Executed in notebook only", "Save fold and system CSVs"),
            ("Paired system win count", "Missing", "Run Phase 5 and report"),
            ("Paired bootstrap 95% interval", "Missing", "Run Phase 5 and replace placeholders"),
            ("System-level A-vs-B plot", "Missing", "Generate from saved system results"),
            ("SMA robustness subset", "Not run", "Complete or state clearly as future work"),
            ("Feature interpretation", "Not run", "Optional grouped ablation or permutation importance"),
            ("Portable data path", "Not complete", "Document retrieval and rebuild process"),
            ("Pinned environment", "Not complete", "Freeze package versions"),
            ("Author/repository metadata", "Placeholders", "Fill and verify before submission"),
        ],
        [3550, 2450, 3360],
        caption="Table A1. Items required to convert this draft into a submission-ready paper",
        source="Source: Repository audit performed for this draft.",
    )


def main():
    fig1, fig2 = make_figures()
    doc = Document()
    configure_document(doc)
    title_page(doc)
    draft_status_page(doc)
    build_paper(doc, fig1, fig2)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
