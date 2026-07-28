from copy import deepcopy
from pathlib import Path
import math

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Solar_PV_Research_Paper_Model_C_Revised.docx"
RESULTS_DIR = ROOT / "model c" / "results"
ASSET_DIR = ROOT / "tmp" / "model_c_paper_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

SUMMARY_PATH = RESULTS_DIR / "phase6_summary.csv"
SYSTEM_PATH = RESULTS_DIR / "phase6_system_results.csv"
PAIRED_PATH = RESULTS_DIR / "phase6_paired_comparison.csv"
FIGURE1_PATH = ASSET_DIR / "figure_1_model_c_design.png"
FIGURE3_PATH = ASSET_DIR / "figure_3_model_c_results.png"

NAVY = "1F4E79"
DEEP_NAVY = "17365D"
BLUE = "3B7FB9"
TEAL = "2A9D8F"
PURPLE = "8C6BBE"
ORANGE = "DC6B35"
GOLD = "C69214"
PALE_BLUE = "EEF4F9"
PALE_TEAL = "E5F4F1"
PALE_GOLD = "FFF4CE"
INK = "202124"
MUTED = "5F6368"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E1E8"
WHITE = "FFFFFF"

FONT_REGULAR = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
FONT_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")


def font(size, bold=False):
    return ImageFont.truetype(
        str(FONT_BOLD if bold else FONT_REGULAR),
        size,
    )


def rgb(hex_color):
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def rounded_box(draw, xy, fill, outline, width=4, radius=18):
    draw.rounded_rectangle(
        xy,
        radius=radius,
        fill=rgb(fill),
        outline=rgb(outline),
        width=width,
    )


def centered_text(draw, xy, text, text_font, fill=INK, spacing=8):
    x0, y0, x1, y1 = xy
    bbox = draw.multiline_textbbox(
        (0, 0),
        text,
        font=text_font,
        spacing=spacing,
        align="center",
    )
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((x0 + x1 - width) / 2, (y0 + y1 - height) / 2),
        text,
        font=text_font,
        fill=rgb(fill),
        spacing=spacing,
        align="center",
    )


def arrow(draw, start, end, color=MUTED, width=5, head=18):
    x0, y0 = start
    x1, y1 = end
    draw.line((x0, y0, x1, y1), fill=rgb(color), width=width)
    angle = math.atan2(y1 - y0, x1 - x0)
    left = (
        x1 - head * math.cos(angle - math.pi / 6),
        y1 - head * math.sin(angle - math.pi / 6),
    )
    right = (
        x1 - head * math.cos(angle + math.pi / 6),
        y1 - head * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=rgb(color))


def make_design_figure():
    image = Image.new("RGB", (2600, 1120), rgb(WHITE))
    draw = ImageDraw.Draw(image)

    draw.text(
        (80, 50),
        "Future-time experimental design",
        font=font(62, bold=True),
        fill=rgb(DEEP_NAVY),
    )
    draw.text(
        (80, 125),
        "System isolation + time isolation + explicit weather-and-time control",
        font=font(34),
        fill=rgb(MUTED),
    )

    top_boxes = [
        ((80, 250, 560, 460), "37 rooftop PV systems\n2021-2023"),
        ((700, 250, 1180, 460), "Five grouped folds\n7-8 systems held out"),
        (
            (1320, 250, 1860, 460),
            "Time boundary\nTrain: through 2022\nTest: 2023",
        ),
        (
            (2000, 250, 2500, 460),
            "System-level metrics\nnRMSE, MAE, MBE",
        ),
    ]
    for box, label in top_boxes:
        rounded_box(draw, box, PALE_BLUE, NAVY)
        centered_text(draw, box, label, font(30, bold=True))
    arrow(draw, (575, 355), (685, 355))
    arrow(draw, (1195, 355), (1305, 355))
    arrow(draw, (1875, 355), (1985, 355))

    model_boxes = [
        (
            (180, 620, 850, 850),
            PALE_BLUE,
            NAVY,
            "Model A\n4 weather features",
        ),
        (
            (965, 620, 1635, 850),
            PALE_TEAL,
            TEAL,
            "Model C\n4 weather + 4 cyclic time features",
        ),
        (
            (1750, 620, 2420, 850),
            PALE_GOLD,
            GOLD,
            "Model B\n4 weather + 5 physics features",
        ),
    ]
    for box, fill, outline, label in model_boxes:
        rounded_box(draw, box, fill, outline, width=5)
        centered_text(draw, box, label, font(32, bold=True))

    arrow(draw, (1500, 475), (515, 605), width=4)
    arrow(draw, (1570, 475), (1300, 605), width=4)
    arrow(draw, (1640, 475), (2085, 605), width=4)

    comparison_box = (455, 930, 2145, 1055)
    rounded_box(draw, comparison_box, LIGHT_GRAY, MUTED, width=3)
    centered_text(
        draw,
        comparison_box,
        "Primary comparison: Model B vs. Model C\n"
        "incremental value of physics beyond ordinary time encodings",
        font(30, bold=True),
        fill=DEEP_NAVY,
    )
    arrow(draw, (1300, 865), (1300, 915), width=4)

    image.save(FIGURE1_PATH)


def make_results_figure(summary, system_results):
    image = Image.new("RGB", (2900, 1320), rgb(WHITE))
    draw = ImageDraw.Draw(image)

    draw.text(
        (70, 35),
        "Future-time model comparison",
        font=font(62, bold=True),
        fill=rgb(INK),
    )
    draw.text(
        (70, 110),
        "Train through 2022; test held-out systems in 2023; n = 37 systems; lower is better",
        font=font(34),
        fill=rgb(MUTED),
    )

    # Left: five-model macro nRMSE comparison.
    draw.text(
        (110, 240),
        "Macro nRMSE by model",
        font=font(42, bold=True),
        fill=rgb(INK),
    )
    plot_left, plot_top, plot_right, plot_bottom = 410, 320, 1640, 1210
    max_value = 0.23
    for tick in [0.00, 0.05, 0.10, 0.15, 0.20]:
        x = plot_left + (plot_right - plot_left) * tick / max_value
        draw.line(
            (x, plot_top, x, plot_bottom),
            fill=rgb(MID_GRAY),
            width=2,
        )
        label = f"{tick:.2f}"
        label_box = draw.textbbox((0, 0), label, font=font(25))
        draw.text(
            (x - (label_box[2] - label_box[0]) / 2, plot_bottom + 15),
            label,
            font=font(25),
            fill=rgb(INK),
        )
    draw.line(
        (plot_left, plot_top, plot_left, plot_bottom),
        fill=rgb(INK),
        width=3,
    )

    indexed = summary.set_index("model")
    model_specs = [
        ("training_mean", "Training mean", "AEB4BA"),
        ("time_solar_only", "Time/solar only", PURPLE),
        ("model_a_weather", "Model A: weather", BLUE),
        ("model_c_weather_time", "Model C: weather + time", TEAL),
        ("model_b_hybrid", "Model B: weather + physics", ORANGE),
    ]
    bar_height = 115
    gap = 45
    for index, (model, label, color) in enumerate(model_specs):
        value = float(indexed.loc[model, "nRMSE_mean"])
        y0 = plot_top + 35 + index * (bar_height + gap)
        y1 = y0 + bar_height
        x1 = plot_left + (plot_right - plot_left) * value / max_value
        draw.rectangle(
            (plot_left, y0, x1, y1),
            fill=rgb(color),
            outline=rgb(INK),
            width=2,
        )
        label_box = draw.textbbox((0, 0), label, font=font(28))
        draw.text(
            (
                plot_left - 22 - (label_box[2] - label_box[0]),
                y0 + (bar_height - (label_box[3] - label_box[1])) / 2,
            ),
            label,
            font=font(28),
            fill=rgb(INK),
        )
        draw.text(
            (x1 + 14, y0 + 35),
            f"{value:.4f}",
            font=font(28),
            fill=rgb(INK),
        )
    axis_label = "Macro nRMSE"
    axis_box = draw.textbbox((0, 0), axis_label, font=font(30))
    draw.text(
        (
            (plot_left + plot_right - (axis_box[2] - axis_box[0])) / 2,
            1275,
        ),
        axis_label,
        font=font(30),
        fill=rgb(INK),
    )

    # Right: Model C vs Model B scatter.
    scatter_left, scatter_top = 1900, 320
    scatter_right, scatter_bottom = 2820, 1210
    draw.text(
        (scatter_left, 240),
        "Physics vs. time-aware control",
        font=font(42, bold=True),
        fill=rgb(INK),
    )
    axis_min, axis_max = 0.09, 0.18
    ticks = [0.10, 0.12, 0.14, 0.16, 0.18]

    def x_coord(value):
        return scatter_left + (
            (value - axis_min)
            / (axis_max - axis_min)
            * (scatter_right - scatter_left)
        )

    def y_coord(value):
        return scatter_bottom - (
            (value - axis_min)
            / (axis_max - axis_min)
            * (scatter_bottom - scatter_top)
        )

    for tick in ticks:
        x = x_coord(tick)
        y = y_coord(tick)
        draw.line(
            (x, scatter_top, x, scatter_bottom),
            fill=rgb(MID_GRAY),
            width=2,
        )
        draw.line(
            (scatter_left, y, scatter_right, y),
            fill=rgb(MID_GRAY),
            width=2,
        )
        tick_label = f"{tick:.2f}"
        tick_box = draw.textbbox((0, 0), tick_label, font=font(24))
        draw.text(
            (x - (tick_box[2] - tick_box[0]) / 2, scatter_bottom + 15),
            tick_label,
            font=font(24),
            fill=rgb(INK),
        )
        draw.text(
            (
                scatter_left - 15 - (tick_box[2] - tick_box[0]),
                y - (tick_box[3] - tick_box[1]) / 2,
            ),
            tick_label,
            font=font(24),
            fill=rgb(INK),
        )
    draw.line(
        (scatter_left, scatter_bottom, scatter_right, scatter_bottom),
        fill=rgb(INK),
        width=3,
    )
    draw.line(
        (scatter_left, scatter_top, scatter_left, scatter_bottom),
        fill=rgb(INK),
        width=3,
    )

    # Dashed equal-error line.
    x0, y0 = x_coord(axis_min), y_coord(axis_min)
    x1, y1 = x_coord(axis_max), y_coord(axis_max)
    segments = 22
    for index in range(0, segments, 2):
        t0 = index / segments
        t1 = min((index + 1) / segments, 1)
        draw.line(
            (
                x0 + (x1 - x0) * t0,
                y0 + (y1 - y0) * t0,
                x0 + (x1 - x0) * t1,
                y0 + (y1 - y0) * t1,
            ),
            fill=rgb(INK),
            width=3,
        )

    wide = system_results.pivot(
        index="station",
        columns="model",
        values="nRMSE",
    )
    no_history = {
        "SQ_Block_P",
        "Shaw_Auditorium",
        "UG_Hall8",
        "UG_Hall9",
    }
    for station, row in wide.iterrows():
        x = x_coord(float(row["model_c_weather_time"]))
        y = y_coord(float(row["model_b_hybrid"]))
        if station in no_history:
            draw.ellipse(
                (x - 12, y - 12, x + 12, y + 12),
                fill=rgb(WHITE),
                outline=rgb(ORANGE),
                width=5,
            )
        else:
            draw.ellipse(
                (x - 10, y - 10, x + 10, y + 10),
                fill=rgb(BLUE),
                outline=rgb(INK),
                width=2,
            )

    x_label = "Model C nRMSE (weather + time)"
    x_label_box = draw.textbbox((0, 0), x_label, font=font(28))
    draw.text(
        (
            (scatter_left + scatter_right - (x_label_box[2] - x_label_box[0]))
            / 2,
            1275,
        ),
        x_label,
        font=font(28),
        fill=rgb(INK),
    )

    y_label = "Model B nRMSE"
    label_layer = Image.new("RGBA", (380, 70), (255, 255, 255, 0))
    label_draw = ImageDraw.Draw(label_layer)
    label_draw.text((0, 0), y_label, font=font(28), fill=rgb(INK))
    rotated = label_layer.rotate(90, expand=True)
    image.paste(
        rotated,
        (
            scatter_left - 120,
            int((scatter_top + scatter_bottom - rotated.height) / 2),
        ),
        rotated,
    )

    # Compact legend.
    legend_y = 345
    draw.ellipse(
        (scatter_left + 25, legend_y, scatter_left + 45, legend_y + 20),
        fill=rgb(BLUE),
        outline=rgb(INK),
        width=2,
    )
    draw.text(
        (scatter_left + 60, legend_y - 5),
        "Held-out system",
        font=font(25),
        fill=rgb(INK),
    )
    draw.ellipse(
        (scatter_left + 25, legend_y + 43, scatter_left + 49, legend_y + 67),
        fill=rgb(WHITE),
        outline=rgb(ORANGE),
        width=4,
    )
    draw.text(
        (scatter_left + 60, legend_y + 40),
        "No pre-2023 history",
        font=font(25),
        fill=rgb(INK),
    )
    draw.text(
        (scatter_left + 25, scatter_bottom - 55),
        "Model B lower on 34 of 37 systems",
        font=font(27, bold=True),
        fill=rgb(DEEP_NAVY),
    )

    image.save(FIGURE3_PATH)


def paragraph_starting(document, prefix):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(prefix)
    ]
    if len(matches) != 1:
        raise ValueError(
            f"Expected one paragraph starting with {prefix!r}; "
            f"found {len(matches)}"
        )
    return matches[0]


def copy_run_properties(source_run, target_run):
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def replace_paragraph(document, prefix, text):
    paragraph = paragraph_starting(document, prefix)
    template_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    run = paragraph.add_run(text)
    copy_run_properties(template_run, run)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(total))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), str(indent_dxa))
    tbl_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_text(cell, text, *, header=False, numeric=False, size=8.9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if numeric or header else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.03
    run = paragraph.add_run(str(text))
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = header
    run.font.color.rgb = RGBColor.from_string(WHITE if header else INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if header:
        shade_cell(cell, NAVY)


def replace_table(document, old_table, headers, rows, widths, numeric_cols=()):
    new_table = document.add_table(rows=1, cols=len(headers))
    new_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    new_table.style = old_table.style or "Table Grid"
    for column_index, header in enumerate(headers):
        set_cell_text(
            new_table.rows[0].cells[column_index],
            header,
            header=True,
            numeric=column_index in numeric_cols,
        )
    set_repeat_table_header(new_table.rows[0])
    for row_index, values in enumerate(rows):
        cells = new_table.add_row().cells
        for column_index, value in enumerate(values):
            set_cell_text(
                cells[column_index],
                value,
                numeric=column_index in numeric_cols,
            )
            if row_index % 2 == 1:
                shade_cell(cells[column_index], LIGHT_GRAY)
    set_table_geometry(new_table, widths)
    set_table_borders(new_table)
    old_table._tbl.addprevious(new_table._tbl)
    old_table._element.getparent().remove(old_table._element)
    return new_table


def replace_image_paragraph(paragraph, image_path, alt_text):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))
    doc_prs = paragraph._p.xpath(".//wp:docPr")
    if doc_prs:
        doc_prs[0].set("descr", alt_text)


def validate_results(summary, system_results, paired):
    assert set(summary["model"]) == {
        "training_mean",
        "time_solar_only",
        "model_a_weather",
        "model_c_weather_time",
        "model_b_hybrid",
    }
    assert len(system_results) == 185
    assert system_results.groupby("model")["station"].nunique().eq(37).all()
    row = paired[
        (paired["comparison"] == "hybrid_vs_time_control")
        & (paired["metric"] == "nRMSE")
    ].iloc[0]
    assert math.isclose(row["reference_macro_mean"], 0.11563173531312076)
    assert math.isclose(row["candidate_macro_mean"], 0.11143215448742531)
    assert row["candidate_better_systems"] == 34
    assert row["relative_reduction_ci95_low"] > 0


def main():
    summary = pd.read_csv(SUMMARY_PATH)
    system_results = pd.read_csv(SYSTEM_PATH)
    paired = pd.read_csv(PAIRED_PATH)
    validate_results(summary, system_results, paired)
    make_design_figure()
    make_results_figure(summary, system_results)

    document = Document(DOCX_PATH)
    document.core_properties.title = (
        "Physics-Derived Features Beyond Time Controls for Future-Time "
        "Prediction on Unseen Rooftop PV Systems"
    )

    replacements = [
        (
            "Physics-Derived Features for Future-Time Prediction",
            "Physics-Derived Features Beyond Time Controls for Future-Time "
            "Prediction on Unseen Rooftop PV Systems",
        ),
        (
            "Data-driven photovoltaic (PV) power models can appear accurate",
            "Data-driven photovoltaic (PV) power models can appear accurate "
            "when training and test data share installations or timestamps, "
            "but new rooftop systems require transfer across both system and "
            "time. This study tests whether physics-derived representations add "
            "predictive value beyond measured weather and ordinary clock and "
            "calendar controls. The HKUST rooftop PV analysis retains 1,370,191 "
            "modeling-ready daylight observations from 37 SolarEdge systems. "
            "The primary experiment trains on pre-2023 observations from other "
            "systems and tests held-out systems during 2023, covering 612,371 "
            "test rows. Weather-only Model A obtains macro nRMSE of 0.1330; "
            "Model C, which adds cyclic hour-of-day and day-of-year features, "
            "reaches 0.1156; and physics-enhanced Model B reaches 0.1114. "
            "Ordinary time features account for 80.6% of the original "
            "Model A-to-B error gap. Model B provides a further 3.6% nRMSE "
            "reduction relative to Model C (paired-bootstrap 95% interval: "
            "2.9%-4.4%), lowers nRMSE on 34 of 37 systems, and lowers MAE on "
            "all 37. Because model inputs are station-invariant at a shared "
            "timestamp, the design estimates a fleet-representative prediction "
            "curve rather than individualized roof behavior. A secondary "
            "same-period comparison is more accurate in absolute terms but "
            "allows exact timestamp-feature overlap. The results support a "
            "modest incremental benefit from the physics representation within "
            "one campus, conditional on contemporaneous measured weather; they "
            "do not establish individualized, day-ahead, or cross-climate "
            "forecasting.",
        ),
        (
            "Keywords:",
            "Keywords: photovoltaics; physics-informed machine learning; "
            "XGBoost; cyclic time controls; domain generalization; cold-start "
            "prediction",
        ),
        (
            "The cold-start setting is especially important",
            "The cold-start setting is especially important for new or "
            "sparsely monitored rooftop systems. A newly commissioned system "
            "has no long history of measured output from which to learn "
            "site-specific behavior. In this study, however, cold-start has a "
            "bounded meaning: the model receives no target-system history and "
            "no system-specific inputs, so it estimates a fleet-representative "
            "normalized-power response to shared environmental conditions. It "
            "does not individualize predictions for roof orientation, shading, "
            "equipment, degradation, or maintenance. Prior research has shown "
            "that shared models can predict individual systems and that "
            "regional models can transfer across installations (Costa, 2022; "
            "Grzebyk et al., 2023). More recent work addresses zero-sample PV "
            "prediction with explicit domain-generalization architectures "
            "(Liu et al., 2025).",
        ),
        (
            "This paper evaluates that hypothesis",
            "This paper evaluates that hypothesis with a controlled future-time "
            "feature comparison. Model A receives four measured weather "
            "variables. Model C adds ordinary cyclic hour-of-day and day-of-year "
            "encodings. Model B instead adds five physics-derived features: "
            "solar position, clear-sky context, and a thermal proxy. All three "
            "XGBoost regressors use the same target, grouped folds, algorithm, "
            "hyperparameters, and random seed. The primary Model B-versus-Model C "
            "difference therefore estimates the incremental predictive value of "
            "the physics representation beyond naive time information under the "
            "chosen protocol. XGBoost is appropriate because tree boosting "
            "captures nonlinear interactions in tabular data and has performed "
            "well in large PV fleets (Chen & Guestrin, 2016; Grzebyk et al., "
            "2023).",
        ),
        (
            "The study makes four contributions.",
            "The study makes four contributions. First, it implements "
            "system-grouped evaluation and a stricter protocol in which both "
            "held-out systems and target timestamps are absent from training. "
            "Second, it introduces an explicit weather-plus-time control so the "
            "physics representation is not credited for merely supplying clock "
            "and calendar information. Third, it documents that all model inputs "
            "are station-invariant at a shared timestamp, which prevents station "
            "fingerprinting but limits the task to a fleet-representative "
            "prediction curve. Fourth, it reports system-level macro metrics and "
            "paired uncertainty so installations with more observations do not "
            "dominate the result.",
        ),
        (
            "The completed same-period analysis reports a 14.0%",
            "The primary future-time result shows that ordinary time controls "
            "reduce macro nRMSE from 0.1330 for weather-only Model A to 0.1156 "
            "for Model C, a 13.1% reduction. Physics-enhanced Model B reaches "
            "0.1114, a further 3.6% reduction relative to Model C with a "
            "paired-bootstrap interval of 2.9%-4.4%. Model B improves nRMSE on "
            "34 of 37 systems and MAE on all 37. The earlier same-period "
            "comparison remains a secondary benchmark: its 0.0919-to-0.0790 "
            "difference measures the value of the complete hybrid feature "
            "bundle under timestamp overlap, not an isolated physics effect.",
        ),
        (
            "The study uses two nested evaluation conditions.",
            "The study uses two nested evaluation conditions. The primary "
            "future-time evaluation withholds complete systems and requires "
            "every test timestamp to occur after every training timestamp. For "
            "each fold, the model is fitted on modeling-ready rows through "
            "December 31, 2022 from the non-held-out systems and evaluated on "
            "2023 rows from the held-out systems. The secondary same-period "
            "evaluation also withholds complete systems but allows training and "
            "test systems to share timestamps. This ordering distinguishes the "
            "main future-time claim from the easier same-period benchmark.",
        ),
        (
            "Two feature sets define the controlled comparison.",
            "Three feature sets define the primary controlled comparison. "
            "Model A uses measured GHI, ambient temperature, wind speed, and "
            "relative humidity. Model C adds sine and cosine encodings of local "
            "hour of day and day of year. Model B uses the four weather "
            "variables plus apparent solar zenith, solar azimuth, clear-sky GHI, "
            "clear-sky index, and a GHI-based Faiman cell-temperature estimate. "
            "The null and alternative hypotheses are:",
        ),
        (
            "H0:",
            "H0: Physics-enhanced Model B does not reduce mean system-level "
            "nRMSE relative to weather-plus-time Model C in the future-time "
            "evaluation.",
        ),
        (
            "H1:",
            "H1: Physics-enhanced Model B reduces mean system-level nRMSE "
            "relative to weather-plus-time Model C in the future-time "
            "evaluation.",
        ),
        (
            "H2:",
            "H2: Any Model B advantage remains broad across held-out systems "
            "and folds when both systems and target timestamps are absent from "
            "training.",
        ),
        (
            "The primary outcome is system-level nRMSE",
            "The primary outcome is system-level nRMSE averaged equally across "
            "held-out systems. Secondary outcomes are normalized MAE and MBE. "
            "Models are compared in paired form because they use the same "
            "systems, folds, timestamps, and eligible rows. The primary contrast "
            "is Model B versus Model C; Model C versus Model A quantifies the "
            "value of ordinary time context, and Model B versus Model A preserves "
            "the total feature-bundle comparison. Neither protocol forecasts "
            "future weather or evaluates operational day-ahead power.",
        ),
        (
            "The daylight filter retained",
            "The daylight filter retained 1,387,682 of 2,756,516 rows. The "
            "maximum retained zenith was 89.9997 degrees. A single campus weather "
            "station supplies the measured weather, and the solar and clear-sky "
            "variables depend only on timestamp and campus location. The analysis "
            "verifies that model inputs do not vary among stations at a shared "
            "timestamp. Consequently, within a fitted fold, held-out systems "
            "observed at the same timestamp receive identical predictions; "
            "separately fitted folds can produce different prediction series. "
            "This structure prevents station fingerprinting through the inputs "
            "but leaves roof-specific variation in normalized output irreducible "
            "under the design.",
        ),
        (
            "4.3 Physics feature construction",
            "4.3 Physics and time feature construction",
        ),
        (
            "Clear-sky GHI is computed",
            "Clear-sky GHI is computed for every retained timestamp with pvlib's "
            "Ineichen model. The clear-sky index is measured GHI divided by "
            "clear-sky GHI and is capped at 2.0. A GHI-based cell-temperature "
            "proxy uses the Faiman model with ambient temperature, wind speed, "
            "u0 = 25.0, and u1 = 6.84. Model C's ordinary time controls are sine "
            "and cosine of local fractional hour on a 24-hour cycle and sine and "
            "cosine of local day of year on a 365.2425-day cycle. These cyclic "
            "encodings preserve adjacency across midnight and the year boundary. "
            "Normalized power is measured power divided by system capacity. No "
            "target-system power values are used to construct weather, time, or "
            "physics features.",
        ),
        (
            "Figure 1.",
            "Figure 1. Primary future-time experimental design. Complete PV "
            "systems are held out, training ends in 2022, and testing uses 2023 "
            "observations. Models A, C, and B use identical folds and XGBoost "
            "settings. The primary Model B-versus-Model C comparison estimates "
            "the incremental value of the physics representation beyond ordinary "
            "time encodings. Source: Study design.",
        ),
        (
            "The principal experiment is a controlled feature ablation.",
            "The principal experiment is a controlled future-time feature "
            "comparison. Models A, C, and B use the same regression algorithm, "
            "objective, hyperparameters, capacity-normalized target, "
            "model-ready rows, grouped folds, and random seed. Model A contains "
            "four measured weather features. Model C adds four cyclic time "
            "features, while Model B adds five physics-derived features. "
            "Plane-of-array irradiance is excluded because a system-level record "
            "may combine arrays with different or uncertain orientations. The "
            "same-period experiment retains the original Model A-versus-Model B "
            "bundle comparison as a secondary benchmark.",
        ),
        (
            "Table 2.",
            "Table 2. Controlled feature sets for the primary future-time "
            "experiment",
        ),
        (
            "Both machine-learning models use XGBoost",
            "All XGBoost models use squared-error regression with the same "
            "configuration. Gradient-boosted trees sequentially correct residual "
            "error while regularization and subsampling reduce overfitting "
            "(Chen & Guestrin, 2016). Tree models do not require feature "
            "standardization, so scaling is disabled. No separate tuning is "
            "performed for Models A, C, and B; tuning independently would "
            "confound the feature comparison.",
        ),
        (
            "Predictions are evaluated at two levels.",
            "Predictions are evaluated at fold and system levels. Headline values "
            "are macro means across the 37 system results, giving each system "
            "equal weight. Because no input identifies a roof, each fitted fold "
            "model emits a common prediction for systems sharing a timestamp. "
            "The estimand is therefore the accuracy of a fleet-representative "
            "environment-to-normalized-power relationship on an unseen roof, not "
            "the accuracy of an individualized roof model.",
        ),
        (
            "Lower nRMSE and MAE indicate",
            "Lower nRMSE and MAE indicate greater accuracy. MBE is signed: "
            "positive values indicate average overprediction and negative values "
            "indicate average underprediction. The primary relative reduction is "
            "100 times the Model C minus Model B error divided by Model C error. "
            "Model C versus Model A measures the time-control gain; Model B "
            "versus Model A reports the total feature-bundle gain.",
        ),
        (
            "The analysis pipeline contains several safeguards.",
            "The pipeline keeps every system in one fold and fits any learned "
            "preprocessing on training rows only. Models share identical group "
            "assignments, and model-ready rows contain no missing inputs or "
            "active exclusion flags. Station ID and system metadata are excluded, "
            "and all supplied features are station-invariant at a shared "
            "timestamp, preventing station fingerprinting through the input "
            "matrix. This does not make every form of leakage impossible: in the "
            "same-period benchmark, other systems contribute labels for the exact "
            "test timestamps and identical feature vectors. The future-time "
            "evaluation removes this information overlap by asserting disjoint "
            "systems, non-overlapping timestamps, and a strict train-before-test "
            "boundary.",
        ),
        (
            "Reproducibility is improved:",
            "Reproducibility is improved: the future-time notebook, Model C "
            "control, split audit, fold results, system results, summaries, and "
            "paired comparisons are preserved together in the project. Fixed "
            "seeds make the folds and bootstrap deterministic. Dependency "
            "versions remain unpinned, and a clean end-to-end rerun has not yet "
            "been documented; final submission should record the environment "
            "and dataset retrieval procedure.",
        ),
        (
            "The uncertainty analysis is paired at the system level.",
            "Uncertainty is calculated with a deterministic paired bootstrap at "
            "the system level. Ten thousand resamples of the 37 paired system "
            "results are drawn with replacement using seed 42. The primary "
            "future-time contrast is Model B minus Model C; the analysis also "
            "reports Model C minus Model A and Model B minus Model A. Resampling "
            "timestamp rows independently would be inappropriate because rows "
            "within a system are temporally dependent. The system bootstrap is "
            "still conditional on an exchangeability approximation: systems "
            "share weather and, within a fold, a common prediction series. Its "
            "intervals should not be interpreted as 37 independent replications "
            "or as cross-location uncertainty.",
        ),
        (
            "The future-time suite contains",
            "The future-time suite contains five models: a training-mean "
            "baseline; a time/solar-only diagnostic using solar zenith, azimuth, "
            "and clear-sky GHI; four-feature weather-only Model A; eight-feature "
            "weather-plus-time Model C; and nine-feature physics-enhanced Model B. "
            "Model C uses the four weather variables plus hour and day-of-year "
            "sine/cosine pairs. The four XGBoost models use the same "
            "configuration, and all five models use identical station folds and "
            "eligible test rows.",
        ),
        (
            "Uncertainty for the future-time evaluation",
            "For the primary Model B-versus-Model C nRMSE comparison, the "
            "system bootstrap reports both the absolute macro difference and the "
            "relative reduction. The same procedure is applied to Model C versus "
            "Model A and Model B versus Model A. Timestamp rows are not resampled "
            "as independent observations, and the shared-campus dependence noted "
            "in Section 5.7 remains a limitation.",
        ),
        (
            "6.2 Physics features reduce average XGBoost error",
            "6.2 Same-period feature bundles reduce average XGBoost error",
        ),
        (
            "The completed same-period analysis reports macro nRMSE",
            "The secondary same-period analysis reports macro nRMSE of 0.0919 "
            "for weather-only Model A and 0.0790 for physics-enhanced Model B. "
            "The absolute difference is -0.0129 of installed capacity, a 14.0% "
            "relative reduction. Macro MAE falls from 0.0614 to 0.0522. Because "
            "Model A contains neither clock variables nor solar geometry, this "
            "comparison estimates the value of the complete weather-plus-physics-"
            "and-time bundle under timestamp overlap. It does not isolate a "
            "physics-specific effect.",
        ),
        (
            "Figure 2.",
            "Figure 2. Secondary same-period feature-bundle comparison. The "
            "left panel shows macro nRMSE across four models. The right panel "
            "compares weather-only Model A with physics-enhanced Model B for each "
            "held-out system; 36 of 37 points favor Model B. Because training "
            "systems supply labels for the same timestamps, the figure does not "
            "separate physics from ordinary time context. Source: Study analysis "
            "outputs.",
        ),
        (
            "Aggregate improvement is not the only relevant result.",
            "Aggregate improvement is not the only relevant same-period result. "
            "Model B's system-level nRMSE standard deviation is 0.0228, compared "
            "with 0.0190 for Model A, and its maximum system error is slightly "
            "worse. The hybrid bundle lowers the macro average and median but "
            "does not improve every system. These results describe the bundle "
            "comparison under overlapping timestamps rather than an isolated "
            "physics mechanism.",
        ),
        (
            "This heterogeneity is scientifically important",
            "This heterogeneity is scientifically important because reliability "
            "on a new system matters in addition to average accuracy. Model B has "
            "lower same-period nRMSE on 36 of 37 systems, but the systems share "
            "weather and feature vectors and should not be treated as independent "
            "confirmations. The scatterplot shows that the bundle gain is broad "
            "within the campus sample while one system favors Model A.",
        ),
        (
            "The observed 14.0% reduction",
            "The same-period 14.0% reduction has a paired-bootstrap interval of "
            "11.3%-16.5%, and the absolute Model B-minus-Model A interval is "
            "-0.0146 to -0.0109. These intervals favor Model B under the "
            "system-resampling procedure, but the shared timestamp inputs and "
            "weather weaken the independence approximation. The result is a "
            "descriptive secondary benchmark, not the paper's primary "
            "physics-beyond-time test.",
        ),
        (
            "6.5 Temporal separation raises absolute error",
            "6.5 The primary future-time evaluation adds an explicit time control",
        ),
        (
            "The station-and-time-disjoint experiment is materially harder",
            "The station-and-time-disjoint experiment is materially harder than "
            "the same-period benchmark. Weather-only macro nRMSE rises from "
            "0.0919 to 0.1330, while hybrid nRMSE rises from 0.0790 to 0.1114. "
            "The mechanism is direct: in the same-period design, the model can "
            "see the exact test-time feature vector during training, labeled with "
            "other systems' normalized output. This is not target-system leakage, "
            "but it gives the model contemporaneous fleet information and makes "
            "the task easier. The 2023 evaluation removes that timestamp overlap.",
        ),
        (
            "Table 5.",
            "Table 5. Primary future-time performance across 37 held-out PV "
            "systems",
        ),
        (
            "Source: Future-time evaluation summary",
            "Source: Model C future-time summary and system-level results. "
            "Training uses pre-2023 rows from other systems; testing uses 2023 "
            "rows from held-out systems.",
        ),
        (
            "The diagnostic models separate",
            "The five-model result separates ordinary time context from the "
            "physics representation. Weather-only Model A obtains macro nRMSE of "
            "0.1330. Model C adds only cyclic hour and day-of-year variables and "
            "reaches 0.1156, a 13.1% reduction with a bootstrap interval of "
            "10.8%-15.2%. Model B reaches 0.1114. In absolute terms, Model C "
            "recovers 80.6% of the original Model A-to-Model B improvement, while "
            "the physics representation accounts for the remaining 19.4%. The "
            "time/solar-only diagnostic remains less accurate at 0.1670 because "
            "it omits measured weather.",
        ),
        (
            "Figure 3.",
            "Figure 3. Primary future-time Model C comparison. The left panel "
            "shows macro nRMSE for all five models when training ends in 2022 "
            "and testing uses held-out systems in 2023. The right panel compares "
            "physics-enhanced Model B with weather-plus-time Model C; Model B is "
            "lower on 34 of 37 systems. Open circles identify four systems with "
            "no eligible pre-2023 history. Source: Model C future-time "
            "system-level results.",
        ),
        (
            "6.6 Physics features improve every future-time held-out system",
            "6.6 Physics adds a modest improvement beyond ordinary time controls",
        ),
        (
            "The system-level paired result is unusually consistent.",
            "Physics-enhanced Model B reduces macro nRMSE from 0.1156 for "
            "weather-plus-time Model C to 0.1114. The absolute difference is "
            "-0.0042 of installed capacity, with a 95% interval from -0.0049 to "
            "-0.0035. The relative reduction is 3.6%, with a 95% interval of "
            "2.9%-4.4%. Model B has lower nRMSE on 34 of 37 systems and in all "
            "five folds. Macro MAE falls from 0.0838 to 0.0790, a 5.7% reduction "
            "with an interval of 4.9%-6.6%, and Model B lowers MAE on all 37 "
            "systems. The original Model A-to-Model B nRMSE reduction remains "
            "16.2%, but most of that total gain is attributable to the time "
            "context supplied by Model C.",
        ),
        (
            "The four systems with no eligible pre-2023 history",
            "The four systems with no eligible pre-2023 history also favor "
            "Model B over Model C. Relative nRMSE reductions are 3.2% for SQ "
            "Block P, 5.4% for Shaw Auditorium, 4.2% for UG Hall 8, and 2.3% "
            "for UG Hall 9. Their average nRMSE is 0.1113 for Model C and 0.1072 "
            "for Model B, a 3.7% reduction. This subgroup supports the bounded "
            "fleet-representative cold-start result, but four correlated campus "
            "systems do not establish individualized or cross-location transfer.",
        ),
        (
            "Bias remains the main weakness.",
            "Bias remains the main weakness. Weather-only Model A has macro MBE "
            "of 0.0670. Model C lowers it to 0.0340 and Model B to 0.0336. The "
            "Model B-minus-Model C bias interval spans -0.0010 to 0.0002, so the "
            "physics representation does not clearly improve mean bias beyond "
            "ordinary time controls. Most of the calibration improvement "
            "relative to Model A is therefore associated with time context, and "
            "both time-aware models retain positive 2023 bias.",
        ),
        (
            "7.1 Why physics-derived features may improve transfer",
            "7.1 What the physics representation adds beyond time controls",
        ),
        (
            "The results are consistent with the idea",
            "The remaining Model B advantage is consistent with physical "
            "transformations providing context that naive time variables do not. "
            "Raw GHI does not state whether irradiance is high or low relative "
            "to the clear-sky expectation. Clear-sky GHI and clear-sky index add "
            "that reference; solar zenith and azimuth represent geometry more "
            "directly than hour and day encodings; and the cell-temperature proxy "
            "combines irradiance, temperature, and wind in a conversion-relevant "
            "form.",
        ),
        (
            "The time/solar-only diagnostic clarifies",
            "Model C materially changes the interpretation. Ordinary time "
            "features explain about 80.6% of the original weather-to-hybrid "
            "nRMSE gap. Model B nevertheless improves a further 3.6% relative to "
            "Model C, with a system-bootstrap interval excluding zero and lower "
            "nRMSE on 34 of 37 systems. The evidence therefore supports a modest "
            "incremental benefit from the complete physics bundle, not the claim "
            "that physics alone caused the full 16.2% reduction. Additional "
            "grouped ablations are needed to identify which physics component "
            "drives the residual gain.",
        ),
        (
            "The results align with prior evidence",
            "The results align with prior evidence that shared models can predict "
            "individual PV systems (Costa, 2022; Grzebyk et al., 2023) and "
            "complement domain-generalization work by Liu et al. (2025). The "
            "present study uses a standard learner and 37 systems from one campus "
            "to compare representations transparently. Model C strengthens that "
            "comparison by distinguishing naive temporal context from the full "
            "physics feature bundle, although the datasets and prediction tasks "
            "remain different from prior studies.",
        ),
        (
            "The comparison with the GHI baseline",
            "The GHI baseline confirms that a simple irradiance-to-power "
            "relationship transfers better than a constant predictor. Weather "
            "Model A improves further, Model C adds substantial ordinary time "
            "context, and Model B provides a smaller additional gain. This "
            "ordering suggests that measured weather, predictable temporal "
            "structure, and physical transformations each contribute under the "
            "future-time protocol.",
        ),
        (
            "In the stricter future-time experiment",
            "In the primary future-time experiment, Model B's nRMSE of 0.1114 "
            "means root-mean-square error of approximately 11.14% of installed "
            "capacity for the average held-out system. For a 100 kW system, that "
            "corresponds to roughly 11.1 kW RMSE under eligible daylight "
            "conditions. Model C's 0.1156 corresponds to about 11.6 kW, so the "
            "incremental physics benefit is approximately 0.42 kW RMSE for a "
            "100 kW system under this normalization. Model B's MAE is about "
            "7.9 kW and its MBE indicates average overprediction of 3.36 kW.",
        ),
        (
            "The same-period evaluation answers",
            "The same-period evaluation asks whether a feature bundle transfers "
            "across systems when other systems supply observations from the same "
            "dates. At a shared timestamp, the feature vector is identical across "
            "systems, so the model has already seen that exact test-time input "
            "paired with other roofs' output. This mechanism explains why "
            "same-period errors are lower. The benchmark remains useful for "
            "historical fleet estimation, but it cannot isolate physics from time "
            "and should not be the primary evidence for future-year transfer.",
        ),
        (
            "Although the hybrid improves nRMSE",
            "Model B improves nRMSE over Model C on 34 of 37 systems, but "
            "absolute error remains heterogeneous. Zone A2, UG Hall 2 2F, and "
            "S H Ho Sports Hall slightly favor Model C in nRMSE, while Model B "
            "lowers MAE on all systems. Because the models receive no roof-specific "
            "features, common clear-sky and temperature proxies cannot represent "
            "orientation, shading, module technology, inverter behavior, "
            "degradation, or maintenance.",
        ),
        (
            "System-level diagnostics should compare",
            "System-level diagnostics should compare error and bias with "
            "available metadata such as capacity, orientation complexity, "
            "commissioning date, and missing-data rate. Such variables are absent "
            "from the current predictors, so between-roof response differences "
            "remain unexplained. The near-equal Model B and Model C bias also "
            "suggests that calibration or transferable system metadata, rather "
            "than additional common timestamp features alone, may be needed to "
            "reduce the positive 2023 shift.",
        ),
        (
            "Several limitations bound the conclusions.",
            "Several limitations bound the conclusions. First, all analyzed "
            "systems share one campus, climate, weather station, and location-"
            "derived solar variables. Within each fitted fold, systems at a "
            "shared timestamp receive identical predictions. This prevents "
            "station fingerprinting and creates a clean fleet-level transfer "
            "test, but it does not evaluate individualized prediction or "
            "geographic, meteorological, and sensor-domain shift. Between-system "
            "variation in normalized output is irreducible without transferable "
            "system-specific inputs.",
        ),
        (
            "Fourth, the 15-minute rows within a system",
            "Fourth, the 15-minute rows within a system are temporally dependent, "
            "and systems share the same weather and, within a fold, the same "
            "prediction series. The 37 system metrics are therefore correlated "
            "targets, not 37 independent confirmations. System-level macro "
            "evaluation and paired system resampling use the correct primary "
            "unit, but the bootstrap's independence approximation is weaker than "
            "it would be across locations and may understate broader uncertainty. "
            "A time-block or two-way system-and-time bootstrap is an appropriate "
            "future sensitivity analysis.",
        ),
        (
            "Fifth, temporal robustness is evaluated",
            "Fifth, temporal robustness is evaluated with one cutoff and one "
            "future year. Four systems have no eligible pre-2023 history and "
            "provide direct evidence for the bounded no-history setting, but "
            "their predictions still follow the same fleet-representative "
            "timestamp relationship. Rolling cutoffs and additional years are "
            "needed to estimate year-to-year variability.",
        ),
        (
            "Seventh, the future-time evaluation result tables",
            "Seventh, the future-time notebook, Model C result tables, and "
            "figures are archived with the experiment, but package versions "
            "remain unpinned and a clean end-to-end rerun has not been "
            "documented. The saved outputs are internally consistent and "
            "reproducible from the station-level CSV; final submission should "
            "record the environment and dataset retrieval procedure.",
        ),
        (
            "Finally, the model configuration was selected",
            "Finally, the shared XGBoost configuration was selected before the "
            "paired comparisons but not through nested tuning. Holding it fixed "
            "is appropriate for comparing Models A, C, and B, yet it does not "
            "show that any feature set is optimally tuned. The findings apply to "
            "one reasonable common configuration.",
        ),
        (
            "The first priority is to move",
            "The first priority is to move from contemporaneous measured weather "
            "to operational weather forecasts and to repeat the future-time "
            "experiment with rolling temporal cutoffs. A block-bootstrap "
            "sensitivity analysis over systems and days or weeks would also "
            "measure uncertainty under shared weather more conservatively.",
        ),
        (
            "A second priority is robustness",
            "A second priority is robustness across hardware and location. The "
            "evaluation should be repeated in an independent climate and with "
            "transferable system metadata such as orientation, shading, module "
            "type, inverter characteristics, and commissioning date. This would "
            "test both cross-location generalization and individualized "
            "cold-start prediction.",
        ),
        (
            "A third direction is mechanism-focused ablation",
            "A third direction is mechanism-focused ablation. Adding solar "
            "geometry, clear-sky normalization, and temperature features in "
            "separate groups would identify which component produces the 3.6% "
            "incremental gain beyond Model C. Calibration should also be "
            "evaluated using training-system data only to address the positive "
            "2023 bias without target-system history.",
        ),
        (
            "This study asks whether a compact set",
            "This study asks whether physics-derived variables improve "
            "future-time PV power prediction beyond measured weather and ordinary "
            "time controls for systems absent from training. The secondary "
            "same-period benchmark gives macro nRMSE of 0.0919 for weather-only "
            "Model A and 0.0790 for Model B, but exact timestamp-feature overlap "
            "makes that comparison easier and prevents physics-specific "
            "attribution.",
        ),
        (
            "The future-time result corresponds",
            "In the primary future-time evaluation, weather-only Model A obtains "
            "macro nRMSE of 0.1330, weather-plus-time Model C obtains 0.1156, and "
            "physics-enhanced Model B obtains 0.1114. Model C explains 80.6% of "
            "the original Model A-to-Model B gap. Model B provides a further "
            "3.6% reduction relative to Model C (95% interval: 2.9%-4.4%), "
            "lowers nRMSE on 34 of 37 systems and all five folds, and lowers MAE "
            "by 5.7% on all 37 systems. Bias is not clearly better than Model C.",
        ),
        (
            "The most defensible conclusion",
            "The most defensible conclusion is bounded: the complete physics "
            "representation provides a modest incremental accuracy benefit beyond "
            "naive time encoding in this campus dataset and controlled XGBoost "
            "pipeline. Most of the original weather-to-hybrid gain comes from "
            "temporal context. Because inputs are station-invariant, the model "
            "estimates a fleet-representative curve rather than individualized "
            "roof behavior; because contemporaneous weather is supplied and all "
            "systems share one climate, the study does not establish day-ahead "
            "or cross-climate forecasting. The shared prediction environment "
            "also means that 37 system-level results should not be interpreted "
            "as independent replications.",
        ),
    ]

    for prefix, text in replacements:
        replace_paragraph(document, prefix, text)

    feature_table = document.tables[1]
    feature_rows = [
        ("Measured GHI", "Yes", "Yes", "Yes", "Current solar resource"),
        (
            "Ambient temperature",
            "Yes",
            "Yes",
            "Yes",
            "Environmental and thermal context",
        ),
        ("Wind speed", "Yes", "Yes", "Yes", "Convective cooling context"),
        (
            "Relative humidity",
            "Yes",
            "Yes",
            "Yes",
            "Atmospheric context",
        ),
        ("Hour sine/cosine", "No", "Yes", "No", "Cyclic time-of-day control"),
        (
            "Day-of-year sine/cosine",
            "No",
            "Yes",
            "No",
            "Cyclic seasonal control",
        ),
        ("Apparent zenith", "No", "No", "Yes", "Solar elevation geometry"),
        ("Solar azimuth", "No", "No", "Yes", "Solar direction geometry"),
        (
            "Clear-sky GHI",
            "No",
            "No",
            "Yes",
            "Time/location irradiance reference",
        ),
        (
            "Clear-sky index",
            "No",
            "No",
            "Yes",
            "Observed irradiance relative to clear sky",
        ),
        (
            "GHI-based cell temperature",
            "No",
            "No",
            "Yes",
            "Approximate thermal state",
        ),
    ]
    replace_table(
        document,
        feature_table,
        ["Feature", "Model A", "Model C", "Model B", "Rationale"],
        feature_rows,
        [2450, 780, 780, 780, 4570],
        numeric_cols=(1, 2, 3),
    )

    future_table = document.tables[4]
    indexed = summary.set_index("model")
    future_models = [
        ("training_mean", "Training mean"),
        ("time_solar_only", "Time/solar-only XGBoost"),
        ("model_a_weather", "Model A: weather-only XGBoost"),
        ("model_c_weather_time", "Model C: weather + time XGBoost"),
        ("model_b_hybrid", "Model B: physics-enhanced XGBoost"),
    ]
    future_rows = []
    for model, label in future_models:
        row = indexed.loc[model]
        future_rows.append(
            (
                label,
                f"{row['nRMSE_mean']:.4f}",
                f"{row['nRMSE_std']:.4f}",
                f"{row['MAE_mean']:.4f}",
                f"{row['MBE_mean']:.4f}",
            )
        )
    replace_table(
        document,
        future_table,
        ["Model", "Macro nRMSE", "SD", "Macro MAE", "Macro MBE"],
        future_rows,
        [3300, 1550, 1150, 1600, 1760],
        numeric_cols=(1, 2, 3, 4),
    )

    image_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.xpath(".//a:blip")
    ]
    if len(image_paragraphs) != 3:
        raise ValueError(f"Expected three figure paragraphs, found {len(image_paragraphs)}")
    replace_image_paragraph(
        image_paragraphs[0],
        FIGURE1_PATH,
        "Future-time design: 37 PV systems are split into five grouped folds; "
        "training ends in 2022 and testing uses 2023. Model A uses weather, "
        "Model C uses weather plus cyclic time, and Model B uses weather plus "
        "physics. The primary comparison is Model B versus Model C.",
    )
    replace_image_paragraph(
        image_paragraphs[2],
        FIGURE3_PATH,
        "Left: macro nRMSE for five future-time models. Right: Model C versus "
        "Model B nRMSE for 37 held-out systems, with four systems lacking "
        "pre-2023 history shown as open circles. Model B is lower on 34 systems.",
    )

    document.save(DOCX_PATH)
    print(f"Saved {DOCX_PATH}")
    print(f"Saved {FIGURE1_PATH}")
    print(f"Saved {FIGURE3_PATH}")


if __name__ == "__main__":
    main()
