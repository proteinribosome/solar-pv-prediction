import base64
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Solar_PV_Research_Paper_Draft.docx"
OUTPUT_DOCX = ROOT / "Solar_PV_Research_Paper_Updated.docx"
NOTEBOOK = ROOT / "code.ipynb"
FIGURE_PATH = ROOT / "tmp" / "paper_assets" / "phase5_main_comparison.png"


def paragraph_starting(document, prefix):
    matches = [p for p in document.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph starting with {prefix!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(document, prefix, text):
    paragraph = paragraph_starting(document, prefix)
    paragraph.clear()
    paragraph.add_run(text)
    return paragraph


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(text)


def extract_phase5_figure():
    notebook = json.loads(NOTEBOOK.read_text())
    image_data = None
    for cell in notebook["cells"]:
        source = "".join(cell.get("source", []))
        if "phase5_main_comparison.png" not in source:
            continue
        for output in cell.get("outputs", []):
            image_data = output.get("data", {}).get("image/png")
            if image_data:
                break
        if image_data:
            break
    if not image_data:
        raise ValueError("The executed Phase 5 comparison figure was not found in code.ipynb")
    if isinstance(image_data, list):
        image_data = "".join(image_data)
    FIGURE_PATH.parent.mkdir(parents=True, exist_ok=True)
    FIGURE_PATH.write_bytes(base64.b64decode(image_data))


def replace_figure_two(document):
    caption = paragraph_starting(document, "Figure 2.")
    paragraphs = document.paragraphs
    caption_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph._p is caption._p
    )
    figure_paragraph = paragraphs[caption_index - 1]
    if not figure_paragraph._p.xpath(".//w:drawing"):
        raise ValueError("Expected Figure 2 image immediately before its caption")

    figure_paragraph.clear()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = figure_paragraph.add_run()
    run.add_picture(str(FIGURE_PATH), width=Inches(6.45))
    doc_prs = figure_paragraph._p.xpath(".//wp:docPr")
    if doc_prs:
        doc_prs[0].set(
            "descr",
            "Left: macro nRMSE for four models. Right: paired system-level nRMSE "
            "for weather-only and hybrid XGBoost across 37 held-out PV systems.",
        )

    caption.clear()
    caption_run = caption.add_run(
        "Figure 2. Phase 5 model comparison. The left panel shows macro nRMSE "
        "across the four models. The right panel compares weather-only Model A "
        "with physics-enhanced Model B for each held-out PV system; points below "
        "the identity line favor Model B. Model B has lower nRMSE on 36 of 37 "
        "systems. Source: Study analysis outputs."
    )
    caption_run.bold = False
    caption_run.italic = False


def set_figure_one_alt_text(document):
    caption = paragraph_starting(document, "Figure 1.")
    paragraphs = document.paragraphs
    caption_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph._p is caption._p
    )
    figure_paragraph = paragraphs[caption_index - 1]
    doc_prs = figure_paragraph._p.xpath(".//wp:docPr")
    if len(doc_prs) != 1:
        raise ValueError("Expected one image immediately before the Figure 1 caption")
    doc_prs[0].set(
        "descr",
        "Experimental design flow: 37 rooftop PV systems are filtered and "
        "quality-controlled, assigned to five grouped folds, evaluated with "
        "system-level metrics, and compared using weather-only Model A and "
        "physics-enhanced Model B.",
    )


def update_completion_table(document):
    candidates = [
        table
        for table in document.tables
        if table.rows
        and table.rows[0].cells[0].text.strip() == "Required item"
    ]
    if len(candidates) != 1:
        raise ValueError(f"Expected one completion table; found {len(candidates)}")
    table = candidates[0]
    updates = {
        "Phase 4 model summaries": (
            "Complete in notebook",
            "Archive fold- and system-level CSVs",
        ),
        "Paired system win count": (
            "Complete: 36 of 37",
            "Retain with the paired comparison",
        ),
        "Paired bootstrap 95% interval": (
            "Complete: 11.3%-16.5%",
            "Retain method, seed, and 10,000 resamples",
        ),
        "System-level A-vs-B plot": (
            "Complete in notebook and paper",
            "Archive the exported PNG with results",
        ),
    }
    for row in table.rows[1:]:
        item = row.cells[0].text.strip()
        if item in updates:
            status, action = updates[item]
            set_cell_text(row.cells[1], status)
            set_cell_text(row.cells[2], action)


def main():
    extract_phase5_figure()
    document = Document(SOURCE_DOCX)

    replacements = [
        (
            "The following items are intentionally unfinished",
            "Phase 5 has now been executed and its paired results are incorporated "
            "in this draft. The following items still require attention before "
            "submission:",
        ),
        (
            "Run Phase 5 and save",
            "Archive the Phase 4 and Phase 5 system-level results, headline "
            "comparison, and exported figure; the values are present in the "
            "completed analysis but have not all been preserved with the project.",
        ),
        (
            "Replace every bracketed [TO COMPLETE]",
            "Record the exact software package versions used in the final analysis.",
        ),
        (
            "Add the system-level Model A versus Model B scatterplot",
            "Reproduce the included system-level Model A versus Model B scatterplot "
            "from archived CSV results during the final clean rerun.",
        ),
        (
            "Evidence status.",
            "Evidence status. Phase 3 results are present in the saved analysis materials. "
            "Phase 4 and Phase 5 were completed; Phase 5 validates four models, "
            "37 systems, and five paired folds, and its output includes the paired "
            "bootstrap results and comparison figure. The Phase 4/5 CSV and PNG "
            "exports are not currently present in the local results folder, so they "
            "should be copied from the execution environment or regenerated before "
            "submission.",
        ),
        (
            "Data-driven photovoltaic (PV) power models often achieve",
            "Data-driven photovoltaic (PV) power models often achieve low error when "
            "their training and test sets contain observations from the same "
            "installation, but this evaluation does not answer a practical cold-start "
            "question: can one model predict the output of a PV system whose historical "
            "power data were never used in training? This study investigates whether "
            "physically derived, system-independent features improve such cross-system "
            "generalization. The analysis uses the open HKUST rooftop PV dataset, which "
            "contains three years of power and meteorological measurements from 60 "
            "campus installations. After restricting the analysis to 37 SolarEdge "
            "systems, localizing timestamps, filtering nighttime observations, "
            "constructing physics features, and applying quality-control rules, "
            "1,370,191 modeling-ready observations remained. Four predictors were "
            "compared under five-fold grouped cross-validation in which entire PV "
            "systems, rather than individual timestamps, were held out: a training-mean "
            "baseline, a calibrated global-horizontal-irradiance baseline, weather-only "
            "XGBoost, and an otherwise identical XGBoost model augmented with solar "
            "zenith, solar azimuth, clear-sky irradiance, clear-sky index, and estimated "
            "cell temperature. Macro normalized root-mean-square errors were 0.2160, "
            "0.1237, 0.0919, and 0.0790, respectively. Relative to weather-only XGBoost, "
            "the physics-enhanced model reduced macro normalized error by 14.0% (paired "
            "bootstrap 95% interval: 11.3%-16.5%) and mean absolute error by 15.0%. It "
            "had lower nRMSE on 36 of 37 held-out systems. These results support the "
            "value of compact physics-derived features for cross-system transfer within "
            "one campus, while cross-climate validation remains necessary before making "
            "a broader generalization claim.",
        ),
        (
            "The executed notebook reports a 14.0% reduction",
            "The completed analysis reports a 14.0% reduction in macro normalized "
            "root-mean-square error (nRMSE) for the hybrid model relative to the "
            "weather-only model. The paired Phase 5 analysis strengthens that "
            "aggregate result: Model B has lower nRMSE on 36 of 37 systems, and the "
            "paired-bootstrap 95% interval for the relative reduction is 11.3%-16.5%. "
            "All systems nevertheless share one campus, climate, and weather station, "
            "so the experiment evaluates cross-system rather than cross-climate "
            "generalization. The paper therefore treats the result as evidence within "
            "a bounded empirical setting, not as proof that the model will transfer to "
            "arbitrary locations.",
        ),
        (
            "Reproducibility is not yet complete",
            "Reproducibility is improved but not yet complete. The workflow still "
            "relies on a temporary cloud-runtime input location, dependency versions "
            "have not been pinned, and some Phase 4 and Phase 5 result exports have "
            "not been preserved with the project. Before submission, the workflow "
            "should use portable inputs, record software versions, document dataset "
            "retrieval, and preserve all non-sensitive derived result tables and figures.",
        ),
        (
            "5.7 Planned statistical comparison",
            "5.7 Paired statistical comparison",
        ),
        (
            "The appropriate uncertainty analysis is paired",
            "The uncertainty analysis is paired at the system level. For each held-out "
            "system, the nRMSE difference is computed as Model B minus Model A. A "
            "deterministic paired bootstrap resamples the 37 systems with replacement "
            "10,000 times using seed 42, recalculates the macro difference and relative "
            "reduction, and uses the 2.5th and 97.5th percentiles as a 95% interval. "
            "Resampling rows would understate uncertainty because adjacent timestamps "
            "from the same system are dependent. The executed comparison gives a mean "
            "nRMSE difference of -0.0129, with a 95% interval from -0.0146 to -0.0109; "
            "the corresponding relative reduction is 14.0%, with a 95% interval from "
            "11.3% to 16.5%. Model B has lower nRMSE on 36 of 37 systems, with no ties.",
        ),
        (
            "The executed Phase 4 notebook output reports",
            "The completed Phase 5 analysis reports macro nRMSE of 0.0919 for "
            "the weather-only model and 0.0790 for the physics-enhanced model. The "
            "absolute difference is -0.0129 of installed capacity, corresponding to a "
            "14.0% relative reduction. Macro MAE falls from 0.0614 to 0.0522, a 15.0% "
            "relative reduction. Mean bias also moves closer to zero, from 0.0010 for "
            "Model A to -0.0004 for Model B.",
        ),
        (
            "Source: Phase 3 CSV files and executed Phase 4 output",
            "Source: Study analysis outputs.",
        ),
        (
            "6.3 The average gain may not be uniform across systems",
            "6.3 The gain is broad but not universal across systems",
        ),
        (
            "This heterogeneity is scientifically important",
            "This heterogeneity is scientifically important because the practical "
            "question concerns reliability on a new system, not only average accuracy. "
            "The paired Phase 5 comparison shows that Model B has lower nRMSE on 36 of "
            "37 systems, with no ties. The scatterplot in Figure 2 places nearly every "
            "system below the identity line, indicating that the macro improvement is "
            "broad rather than driven by a small subset. One system nevertheless favors "
            "Model A, consistent with Model B's slightly worse maximum system error.",
        ),
        (
            "6.4 Statistical uncertainty remains pending",
            "6.4 Paired uncertainty supports the within-campus improvement",
        ),
        (
            "The observed 14.0% reduction is a point estimate",
            "The observed 14.0% reduction is a point estimate over 37 systems. In the "
            "10,000-resample paired bootstrap, the 95% interval for the relative nRMSE "
            "reduction is 11.3%-16.5%, and the interval for the absolute Model B minus "
            "Model A difference is -0.0146 to -0.0109. Both intervals favor Model B. "
            "Because the resampling unit is the held-out PV system, this result supports "
            "a stable improvement across systems represented by the campus sample. It "
            "does not establish causality or guarantee transfer to different climates, "
            "hardware, or data-generating processes.",
        ),
        (
            "Fifth, the current result is not fully reproducible",
            "Fifth, the current result is not yet fully reproducible from a clean "
            "environment. The analysis inputs are available, but the workflow still "
            "uses a cloud-runtime-specific location, dependency versions are unpinned, "
            "and some derived Phase 4/5 results and figure exports have not been "
            "preserved with the project. The saved analysis outputs support the reported "
            "findings, but a clean rerun from portable, documented inputs should be "
            "completed before submission.",
        ),
        (
            "The first priority is to complete the existing analysis",
            "The first priority is to archive and reproduce the completed analysis "
            "rather than add model complexity. The Phase 4 and Phase 5 fold- and "
            "system-level tables, headline comparison, bootstrap results, and paired "
            "scatterplot should be saved in the project results folder. The pipeline "
            "should then be rerun from a portable dataset path with pinned versions and "
            "fixed seeds.",
        ),
        (
            "These results support the practical value",
            "These results support the practical value of solar position, clear-sky "
            "context, and an approximate cell-temperature feature for cross-system "
            "transfer within one campus. Model B lowers nRMSE on 36 of 37 held-out "
            "systems, and the paired bootstrap places the relative improvement between "
            "11.3% and 16.5%. The hybrid's larger error spread and slightly worse maximum "
            "error nevertheless show that the gain is not universal.",
        ),
        (
            "The most defensible current conclusion is therefore bounded",
            "The most defensible conclusion is therefore bounded: physics-derived "
            "features improve held-out-system prediction in this specific dataset and "
            "controlled XGBoost pipeline, with a broad and statistically stable paired "
            "advantage across the sampled systems. Testing another climate and hardware "
            "population remains necessary before claiming general transfer to unseen "
            "rooftop PV systems.",
        ),
        (
            "OpenAI Codex was used on July 21, 2026",
            "OpenAI Codex was used on July 21, 2026 to create an initial paper structure, "
            "draft prose, format the Word document, and identify incomplete evidence. On "
            "July 25, 2026, it was used to update the draft from the completed Phase 5 "
            "analysis, incorporate the paired statistics and comparison figure, and "
            "check the revised document layout. On July 26, 2026, it was used to remove "
            "internal implementation references, revise reproducibility language, and "
            "audit submission readiness. The tool was given the course "
            "syllabus, the local project notebook, saved result tables, and primary "
            "literature metadata. It did not independently rerun the full raw-data and "
            "model-training pipeline. The student author is responsible for checking "
            "every claim, revising the text, documenting the interaction as required by "
            "the Pioneer Research Program, and ensuring that the final paper represents "
            "the student's own reasoning and research. Suggested citation: OpenAI. "
            "(2026). Codex [Large language model]. https://openai.com/codex/",
        ),
    ]

    for prefix, text in replacements:
        replace_paragraph(document, prefix, text)

    set_figure_one_alt_text(document)
    replace_figure_two(document)
    update_completion_table(document)

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
