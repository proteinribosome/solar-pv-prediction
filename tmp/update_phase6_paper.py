from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Solar_PV_Research_Paper.docx"
OUTPUT_DOCX = ROOT / "Solar_PV_Research_Paper_Phase6.docx"
RESULTS_DIR = ROOT / "future_forecasting_experiment" / "results"
FIGURE_PATH = RESULTS_DIR / "phase6_main_comparison_paper.png"
SUMMARY_PATH = RESULTS_DIR / "phase6_summary.csv"
PAIRED_PATH = RESULTS_DIR / "phase6_paired_comparison.csv"

NAVY = "1F4E79"
DEEP_NAVY = "17365D"
INK = "202124"
MUTED = "5F6368"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E1E8"
WHITE = "FFFFFF"


def paragraph_starting(document, prefix):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(prefix)
    ]
    if len(matches) != 1:
        raise ValueError(
            f"Expected one paragraph starting with {prefix!r}; "
            f"found {len(matches)}"
        )
    return matches[0]


def copy_run_properties(source_run, target_run):
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def replace_paragraph(document, prefix, text):
    paragraph = paragraph_starting(document, prefix)
    template_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    run = paragraph.add_run(text)
    copy_run_properties(template_run, run)
    return paragraph


def insert_paragraph_before(document, anchor, text, style=None):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(total))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), str(indent_dxa))
    tbl_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_text(cell, text, *, header=False, numeric=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if numeric or header else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(9.2)
    run.bold = header
    run.font.color.rgb = RGBColor.from_string(WHITE if header else INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if header:
        shade_cell(cell, NAVY)


def insert_results_table(document, anchor, summary):
    caption = document.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = True
    caption_run = caption.add_run(
        "Table 5. Future-time performance across 37 held-out PV systems"
    )
    caption_run.bold = True
    caption_run.font.color.rgb = RGBColor.from_string(DEEP_NAVY)
    anchor._p.addprevious(caption._p)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    headers = ["Model", "Macro nRMSE", "SD", "Macro MAE", "Macro MBE"]
    for index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            header,
            header=True,
            numeric=index > 0,
        )
    set_repeat_table_header(table.rows[0])

    labels = {
        "training_mean": "Training mean",
        "time_solar_only": "Time/solar-only XGBoost",
        "model_a_weather": "Model A: weather-only XGBoost",
        "model_b_hybrid": "Model B: physics-enhanced XGBoost",
    }
    order = [
        "training_mean",
        "time_solar_only",
        "model_a_weather",
        "model_b_hybrid",
    ]
    indexed = summary.set_index("model")
    for row_index, model in enumerate(order):
        result = indexed.loc[model]
        values = [
            labels[model],
            f"{result['nRMSE_mean']:.4f}",
            f"{result['nRMSE_std']:.4f}",
            f"{result['MAE_mean']:.4f}",
            f"{result['MBE_mean']:.4f}",
        ]
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            set_cell_text(cells[column_index], value, numeric=column_index > 0)
            if row_index % 2 == 1:
                shade_cell(cells[column_index], LIGHT_GRAY)

    set_table_geometry(table, [3400, 1550, 1150, 1550, 1710])
    set_table_borders(table)
    anchor._p.addprevious(table._tbl)

    source = document.add_paragraph(style="Table Source")
    source_run = source.add_run(
        "Source: Phase 6 summary and system-level results. Training uses "
        "pre-2023 rows from other systems; testing uses 2023 rows from "
        "held-out systems."
    )
    source_run.italic = True
    source_run.font.color.rgb = RGBColor.from_string(MUTED)
    anchor._p.addprevious(source._p)


def insert_figure(document, anchor):
    figure_paragraph = document.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.keep_with_next = True
    run = figure_paragraph.add_run()
    run.add_picture(str(FIGURE_PATH), width=Inches(6.35))
    doc_prs = figure_paragraph._p.xpath(".//wp:docPr")
    if doc_prs:
        doc_prs[0].set(
            "descr",
            "Left: macro nRMSE for four future-time prediction models. "
            "Right: weather-only versus hybrid nRMSE for 37 held-out PV "
            "systems, with four systems lacking pre-2023 history shown as "
            "open circles. Every point lies below the equal-error line.",
        )
    anchor._p.addprevious(figure_paragraph._p)

    caption = document.add_paragraph(style="Caption")
    caption_run = caption.add_run(
        "Figure 3. Phase 6 future-time comparison. The left panel shows "
        "macro nRMSE when training ends in 2022 and testing uses 2023 "
        "observations from held-out systems. The right panel compares "
        "weather-only and hybrid nRMSE for each system; all 37 points favor "
        "the hybrid. Open circles identify four systems with no eligible "
        "pre-2023 history. Source: Phase 6 system-level results."
    )
    caption_run.bold = False
    caption_run.italic = False
    anchor._p.addprevious(caption._p)


def validate_sources(summary, paired):
    if set(summary["model"]) != {
        "training_mean",
        "time_solar_only",
        "model_a_weather",
        "model_b_hybrid",
    }:
        raise ValueError("Unexpected Phase 6 model set")
    nrmse = paired.loc[paired["metric"].eq("nRMSE")].iloc[0]
    mae = paired.loc[paired["metric"].eq("MAE")].iloc[0]
    checks = [
        (nrmse["n_systems"], 37),
        (nrmse["candidate_better_systems"], 37),
        (nrmse["relative_error_reduction_pct"], 16.246531571177783),
        (mae["relative_error_reduction_pct"], 19.8546810408544),
    ]
    for actual, expected in checks:
        if not abs(float(actual) - float(expected)) < 1e-10:
            raise ValueError(f"Phase 6 source check failed: {actual} != {expected}")
    if not FIGURE_PATH.exists():
        raise FileNotFoundError(FIGURE_PATH)


def main():
    summary = pd.read_csv(SUMMARY_PATH)
    paired = pd.read_csv(PAIRED_PATH)
    validate_sources(summary, paired)

    document = Document(SOURCE_DOCX)
    document.core_properties.title = (
        "Physics-Derived Features for Future-Time Prediction on Unseen "
        "Rooftop PV Systems"
    )

    replacements = [
        (
            "Physics-Derived Features and Cross-System Generalization",
            "Physics-Derived Features for Future-Time Prediction on Unseen "
            "Rooftop PV Systems",
        ),
        (
            "Data-driven photovoltaic (PV) power models often achieve",
            "Data-driven photovoltaic (PV) power models can appear accurate "
            "when training and test data share installations or timestamps, "
            "but new rooftop systems require transfer across both system and "
            "time. This study tests whether physically derived, "
            "system-independent features improve prediction for unseen PV "
            "systems under two evaluation protocols. The analysis uses the "
            "HKUST rooftop PV dataset and retains 1,370,191 modeling-ready "
            "daylight observations from 37 SolarEdge systems. Weather-only "
            "XGBoost is compared with an otherwise identical hybrid model "
            "augmented with solar zenith, solar azimuth, clear-sky "
            "irradiance, clear-sky index, and estimated cell temperature. "
            "In the same-period leave-systems-out evaluation, macro nRMSE "
            "falls from 0.0919 to 0.0790, a 14.0% reduction. A stricter "
            "future-time experiment trains only on pre-2023 observations "
            "from other systems and tests held-out systems during 2023, "
            "covering 612,371 test rows. In this setting, a time/solar-only "
            "model, weather-only XGBoost, and the hybrid obtain macro nRMSE "
            "of 0.1670, 0.1330, and 0.1114, respectively. The hybrid reduces "
            "nRMSE by 16.2% relative to weather-only XGBoost (paired "
            "bootstrap 95% interval: 14.2%-18.2%) and has lower error on all "
            "37 systems, including four with no eligible pre-2023 history. "
            "Absolute error is higher than in the same-period experiment and "
            "the hybrid overpredicts 2023 output by 3.36% of capacity on "
            "average. The results support physics-derived features for "
            "future-time prediction on unseen systems within one campus, "
            "conditional on contemporaneous measured weather; they do not "
            "establish day-ahead or cross-climate forecasting.",
        ),
        (
            "The study makes three contributions.",
            "The study makes four contributions. First, it implements a "
            "system-grouped comparison in which all observations from a test "
            "system are excluded from training. Second, it adds a stricter "
            "future-time protocol in which test systems and test timestamps "
            "are both unseen. Third, it compares a time/solar-only diagnostic, "
            "weather-only XGBoost, and a physics-enhanced hybrid to determine "
            "whether performance can be explained by the typical solar-time "
            "curve alone. Fourth, it reports system-level macro metrics and "
            "paired uncertainty so installations with more observations do "
            "not dominate the conclusion.",
        ),
        (
            "The completed analysis reports a 14.0% reduction",
            "The completed same-period analysis reports a 14.0% reduction in "
            "macro nRMSE for the hybrid relative to weather-only XGBoost. The "
            "future-time analysis is harder in absolute terms, but the hybrid "
            "still reduces macro nRMSE from 0.1330 to 0.1114, a 16.2% relative "
            "reduction with a paired-bootstrap 95% interval of 14.2%-18.2%. "
            "It improves all 37 systems, while a time/solar-only model remains "
            "substantially less accurate at 0.1670. All systems nevertheless "
            "share one campus, climate, and weather station, and the future "
            "models receive contemporaneous measured weather. The result "
            "therefore supports bounded future-time, cross-system prediction "
            "rather than day-ahead or cross-climate forecasting.",
        ),
        (
            "The central evaluation condition is zero historical power data",
            "The study uses two nested evaluation conditions. Phase 5 tests "
            "zero historical power data from a held-out system while allowing "
            "training and test systems to share timestamps. Phase 6 retains "
            "the same system isolation and additionally requires every test "
            "timestamp to occur after every training timestamp. For each "
            "Phase 6 fold, the regressor is fitted using modeling-ready rows "
            "through December 31, 2022 from the non-held-out systems and is "
            "evaluated on 2023 rows from the held-out systems. This design "
            "distinguishes same-period cross-system transfer from future-time "
            "prediction on an unseen system.",
        ),
        (
            "The primary outcome is system-level nRMSE",
            "The primary outcome is system-level nRMSE averaged equally across "
            "held-out systems. Secondary outcomes are normalized mean absolute "
            "error (MAE) and mean bias error (MBE). Models are compared in "
            "paired form because they use the same systems, folds, timestamps, "
            "and eligible rows. Phase 5 estimates same-period cross-system "
            "transfer; Phase 6 estimates future-time transfer to unseen systems "
            "conditional on contemporaneous measured weather. Neither protocol "
            "forecasts future weather or evaluates operational day-ahead power.",
        ),
        (
            "Figure 1. Experimental design.",
            "Figure 1. Same-period leave-systems-out experimental design. All "
            "timestamps from each held-out PV system remain outside the training "
            "set. Models A and B use identical folds and XGBoost settings, so "
            "their paired error difference isolates the added physics-derived "
            "features. Phase 6 adds the non-overlapping time boundary described "
            "in Section 5.8. Source: Study design.",
        ),
        (
            "The analysis pipeline contains several safeguards.",
            "The analysis pipeline contains several safeguards. All rows from a "
            "system remain in one fold, and preprocessing that can learn from "
            "data is fitted on training rows only. Models A and B use identical "
            "group assignments. The cleaned table retains time-zone information "
            "and verifies unique system-timestamp rows. Model-ready rows contain "
            "no missing model variables or active exclusion flags. Phase 6 adds "
            "explicit assertions that training and test systems are disjoint, "
            "that the latest training timestamp precedes the earliest test "
            "timestamp, and that no timestamp appears in both partitions. Fixed "
            "random seeds support deterministic reruns.",
        ),
        (
            "Reproducibility is improved but not yet complete.",
            "Reproducibility is improved: the Phase 6 notebook uses a portable "
            "experiment-relative data path, and its fold, system, summary, "
            "paired-bootstrap, and figure outputs are preserved together in "
            "the project. Dependency versions remain unpinned, and the complete "
            "pipeline has not yet been reproduced from a clean environment. "
            "Before submission, dataset retrieval and software versions should "
            "be documented and a clean rerun should confirm every saved value.",
        ),
        (
            "The average improvement is consistent with the idea",
            "The results are consistent with the idea that physical "
            "transformations provide a more stable representation across "
            "systems and years. Raw GHI indicates the instantaneous resource, "
            "but it does not directly state whether that irradiance is high or "
            "low relative to the current solar position. Clear-sky GHI and "
            "clear-sky index add that context. Solar zenith and azimuth encode "
            "deterministic daily and seasonal geometry, while the "
            "cell-temperature proxy combines irradiance, ambient temperature, "
            "and wind in a form more closely related to PV conversion.",
        ),
        (
            "Tree ensembles can in principle learn interactions",
            "The time/solar-only diagnostic clarifies that the hybrid is not "
            "merely looking up a typical percentage for a timestamp. Solar "
            "geometry alone improves macro nRMSE from 0.2095 to 0.1670, so the "
            "daily and seasonal curve is useful. However, the full hybrid "
            "reaches 0.1114 and improves every held-out system despite all 2023 "
            "timestamps being absent from training. The five-feature bundle "
            "therefore contributes transferable weather and physical context "
            "beyond a recurring solar-time pattern, although grouped ablation "
            "is still needed to isolate each feature's mechanism.",
        ),
        (
            "An nRMSE of 0.0790 means",
            "In the stricter Phase 6 experiment, an nRMSE of 0.1114 means that "
            "root-mean-square prediction error is approximately 11.14% of "
            "installed capacity for the average held-out system. For a 100 kW "
            "system, this corresponds to roughly 11.1 kW RMSE under eligible "
            "daylight conditions. Macro MAE is 0.0790, or about 7.9 kW for a "
            "100 kW system, and MBE of 0.0336 corresponds to average "
            "overprediction of approximately 3.36 kW. System-specific errors "
            "can be substantially higher or lower.",
        ),
        (
            "That interpretation should not be confused",
            "That interpretation should not be confused with day-ahead "
            "operational forecasting. Phase 6 predicts later timestamps, but "
            "the models receive measured contemporaneous irradiance and "
            "weather at those timestamps. Replacing those inputs with weather "
            "forecasts would introduce additional error and require a defined "
            "forecast horizon. The present model is most directly relevant to "
            "conditional prediction, performance benchmarking, and "
            "fault-screening applications.",
        ),
        (
            "7.4 Why some systems may not improve",
            "7.5 Why future-time error and bias remain system-dependent",
        ),
        (
            "The slightly worse maximum error for Model B",
            "Although the hybrid improves nRMSE on all 37 systems in Phase 6, "
            "absolute error remains heterogeneous. UG Hall 2 2F, Zone A2, "
            "S H Ho Sports Hall, and Zone J2 have the largest hybrid nRMSE, "
            "and each shows substantial positive bias. Rooftops differ in "
            "orientation, shading, module technology, inverter behavior, "
            "commissioning date, degradation, and maintenance. Common "
            "clear-sky and temperature proxies cannot represent all of these "
            "system-specific factors.",
        ),
        (
            "System-level diagnostics should therefore compare",
            "System-level diagnostics should compare error and bias against "
            "available metadata such as capacity, optimizer type, orientation "
            "complexity, commissioning date, and missing-data rate. The hybrid "
            "reduces absolute bias relative to weather-only XGBoost on all 37 "
            "systems, but it still overpredicts on 35. A training-period "
            "calibration that is evaluated without using target-system history, "
            "or richer transferable metadata, may reduce this temporal bias.",
        ),
        (
            "Second, the analysis is conditional power prediction",
            "Second, both evaluations are conditional power prediction rather "
            "than complete operational forecasts. Phase 6 moves the target "
            "timestamps into a later year, but contemporaneous measured weather "
            "is still supplied to the models. Numerical weather-prediction "
            "inputs would add forecast error and require an explicit horizon. "
            "Future-time conditional prediction is therefore more accurate "
            "terminology than day-ahead forecasting.",
        ),
        (
            "Fifth, the current result is not yet fully reproducible",
            "Fifth, temporal robustness is evaluated with one cutoff and one "
            "future calendar year. Weather regimes, degradation, and system "
            "availability may differ in other years. Four systems have no "
            "eligible pre-2023 history and provide useful cold-start evidence, "
            "but their average hybrid improvement is smaller than for the "
            "remaining systems. Rolling cutoffs and additional years are needed "
            "to estimate year-to-year variability.",
        ),
        (
            "The first priority is to archive and reproduce",
            "The first priority is to move from contemporaneous measured weather "
            "to operationally available weather forecasts and to repeat the "
            "station-and-time-disjoint experiment with rolling temporal cutoffs. "
            "This would establish explicit forecast horizons, quantify weather-"
            "forecast error, and show whether the 2023 result persists across "
            "different years.",
        ),
        (
            "A second priority is robustness across hardware and location.",
            "A second priority is robustness across hardware and location. "
            "The grouped and future-time evaluations should be repeated in an "
            "independent hardware population and on a public dataset from a "
            "different climate, such as Alice Springs. Such tests would show "
            "whether the present within-campus result extends across equipment, "
            "weather sensors, and climate domains.",
        ),
        (
            "A third direction is mechanism-focused ablation.",
            "A third direction is mechanism-focused ablation and calibration. "
            "Adding solar geometry, clear-sky normalization, and temperature "
            "features in groups would identify which representation transfers "
            "across time. Calibration methods should also be evaluated using "
            "training-system data only to address the positive 2023 bias without "
            "introducing target-system history.",
        ),
        (
            "Finally, physics-as-features can be compared",
            "Finally, physics-as-features can be compared with physics-guided "
            "residual correction. In the latter design, a physical model "
            "produces an initial prediction and machine learning models its "
            "residual error. Comparing these strategies under the same "
            "station-and-time-disjoint folds would test whether physical "
            "structure is more transferable as an input representation or as "
            "a baseline model.",
        ),
        (
            "This study asks whether a compact set",
            "This study asks whether a compact set of physics-derived variables "
            "improves PV power prediction for rooftop systems that are absent "
            "from training, including at future timestamps. In the same-period "
            "leave-systems-out evaluation, weather-only and hybrid XGBoost "
            "achieve macro nRMSE of 0.0919 and 0.0790. In the stricter "
            "future-time experiment, training ends in 2022 and held-out systems "
            "are tested during 2023; macro nRMSE is 0.1330 for weather-only "
            "XGBoost and 0.1114 for the hybrid.",
        ),
        (
            "These results support the practical value",
            "The Phase 6 difference corresponds to a 16.2% reduction in macro "
            "nRMSE and a 19.9% reduction in macro MAE. The paired bootstrap "
            "places the relative nRMSE improvement between 14.2% and 18.2%, "
            "and the hybrid lowers nRMSE on all 37 held-out systems, including "
            "four with no eligible pre-2023 history. The time/solar-only model "
            "remains materially less accurate, showing that the result cannot "
            "be explained by the typical daily and seasonal output curve alone.",
        ),
        (
            "The most defensible conclusion is therefore bounded",
            "The most defensible conclusion is bounded but stronger than the "
            "same-period result alone: physics-derived features improve "
            "future-time prediction on unseen rooftop PV systems in this "
            "specific campus dataset and controlled XGBoost pipeline. Absolute "
            "error rises under temporal separation and the hybrid retains "
            "positive 2023 bias. Because contemporaneous measured weather is "
            "supplied and all systems share one climate and weather station, "
            "the study does not establish day-ahead or cross-climate "
            "forecasting.",
        ),
    ]

    for prefix, replacement in replacements:
        replace_paragraph(document, prefix, replacement)

    primary_outcome = paragraph_starting(
        document, "The primary outcome is system-level nRMSE"
    )
    insert_paragraph_before(
        document,
        primary_outcome,
        "H2: The physics-enhanced model retains lower system-level nRMSE "
        "when both the held-out systems and their target timestamps are "
        "absent from training.",
        style="List Bullet",
    )

    results_heading = paragraph_starting(document, "6. Experimental Results")
    insert_paragraph_before(
        document,
        results_heading,
        "5.8 Future-time prediction on unseen systems",
        style="Heading 2",
    )
    insert_paragraph_before(
        document,
        results_heading,
        "Phase 6 assigns every system to one of five deterministic held-out "
        "folds. Training uses only modeling-ready observations through "
        "December 31, 2022 from the non-held-out systems, while testing uses "
        "January 1 through December 31, 2023 observations from the held-out "
        "systems. Each fold contains seven or eight test systems and 117,255 "
        "to 133,631 test rows. Because four systems begin in 2023, the actual "
        "number of systems contributing historical training rows ranges from "
        "26 to 28. Across folds, every one of the 37 systems is tested exactly "
        "once, totaling 612,371 future observations.",
    )
    insert_paragraph_before(
        document,
        results_heading,
        "The future-time suite contains a training-mean baseline, a "
        "time/solar-only XGBoost model using solar zenith, azimuth, and "
        "clear-sky GHI, the four-feature weather-only Model A, and the "
        "nine-feature hybrid Model B. The time/solar-only diagnostic tests "
        "whether a typical daily and seasonal output curve can explain "
        "performance without measured weather. The three XGBoost models use "
        "the same configuration, and all four models use identical station "
        "folds and eligible test rows.",
    )
    insert_paragraph_before(
        document,
        results_heading,
        "Phase 6 uncertainty is calculated with the same system-level paired "
        "bootstrap principle as Phase 5. Ten thousand resamples of the 37 "
        "systems are drawn with replacement using seed 42. For nRMSE, the "
        "reported interval is computed for both the absolute hybrid-minus-"
        "weather difference and the relative reduction. Timestamp rows are "
        "not resampled as independent observations.",
    )

    discussion_heading = paragraph_starting(document, "7. Discussion")
    insert_paragraph_before(
        document,
        discussion_heading,
        "6.6 Temporal separation raises absolute error but preserves the "
        "hybrid advantage",
        style="Heading 2",
    )
    insert_paragraph_before(
        document,
        discussion_heading,
        "The station-and-time-disjoint experiment is materially harder than "
        "the same-period evaluation. Weather-only macro nRMSE rises from "
        "0.0919 to 0.1330, a 44.7% increase, while hybrid nRMSE rises from "
        "0.0790 to 0.1114, a 41.0% increase. The overlapping-time protocol "
        "therefore understates future-year error, but both experiments retain "
        "the same model ordering and a substantial hybrid advantage.",
    )
    insert_results_table(document, discussion_heading, summary)
    insert_paragraph_before(
        document,
        discussion_heading,
        "The diagnostic models separate a typical solar-time curve from "
        "measured weather and physics context. Time/solar-only XGBoost lowers "
        "macro nRMSE from 0.2095 for the training mean to 0.1670, a 20.3% "
        "reduction. Weather-only XGBoost reaches 0.1330 and outperforms the "
        "time/solar-only model on 32 of 37 systems. The hybrid reaches 0.1114, "
        "a 33.3% reduction relative to time/solar-only, and outperforms it on "
        "all 37 systems. This pattern shows that recurring solar geometry "
        "matters but does not explain the hybrid result.",
    )
    insert_figure(document, discussion_heading)
    insert_paragraph_before(
        document,
        discussion_heading,
        "6.7 Physics features improve every future-time held-out system",
        style="Heading 2",
    )
    insert_paragraph_before(
        document,
        discussion_heading,
        "The system-level paired result is unusually consistent. Hybrid "
        "nRMSE is lower than weather-only nRMSE on all 37 systems, with no "
        "ties. The mean absolute difference is -0.0216 of installed capacity "
        "(95% interval: -0.0251 to -0.0182), and the 16.2% relative reduction "
        "has a 95% interval of 14.2%-18.2%. Macro MAE falls from 0.0986 to "
        "0.0790, a 19.9% reduction with a paired-bootstrap interval of "
        "18.0%-21.6%. All five folds favor the hybrid.",
    )
    insert_paragraph_before(
        document,
        discussion_heading,
        "The four systems with no eligible pre-2023 history also favor the "
        "hybrid. Relative nRMSE reductions are 16.8% for SQ Block P, 13.5% "
        "for Shaw Auditorium, 4.3% for UG Hall 8, and 1.8% for UG Hall 9. "
        "Their average hybrid nRMSE is 0.1072 versus 0.1183 for weather-only "
        "XGBoost, a 9.4% reduction. This subgroup provides direct cold-start "
        "evidence, although its smaller and more variable benefit should not "
        "be generalized beyond four systems.",
    )
    insert_paragraph_before(
        document,
        discussion_heading,
        "Bias remains the main weakness. Weather-only XGBoost overpredicts "
        "all 37 systems and has macro MBE of 0.0670. The hybrid reduces "
        "absolute bias on every system and lowers macro MBE to 0.0336, but "
        "still overpredicts on 35 systems. Thus, the physics features improve "
        "both error and calibration without fully correcting the shift "
        "between pre-2023 training data and 2023 outcomes.",
    )

    heading_75 = paragraph_starting(
        document, "7.5 Why future-time error and bias remain"
    )
    insert_paragraph_before(
        document,
        heading_75,
        "7.4 Same-period evaluation is optimistic but directionally reliable",
        style="Heading 2",
    )
    insert_paragraph_before(
        document,
        heading_75,
        "Phase 5 answers whether a model transfers across systems when other "
        "systems supply training observations from the same date range. That "
        "question remains useful for fleet-level estimation and historical "
        "benchmarking, but it is easier than predicting a later year. Phase 6 "
        "shows that temporal overlap lowers absolute error by roughly 41%-45%. "
        "However, the hybrid advantage persists and grows slightly in relative "
        "terms, from 14.0% to 16.2%. The earlier result was therefore "
        "optimistic about absolute accuracy but directionally reliable about "
        "the value of physics-derived features.",
    )

    model_config_paragraph = paragraph_starting(
        document, "Finally, the model configuration was selected"
    )
    insert_paragraph_before(
        document,
        model_config_paragraph,
        "Sixth, the reported population is restricted to modeling-ready "
        "daylight observations. Quality-control rules exclude missing values, "
        "physically impossible ranges, stuck measurements, and likely daytime "
        "fault zeros; some exclusions therefore use observed power. The "
        "reported errors describe expected production during retained "
        "operation and should not be interpreted as performance for fault or "
        "outage detection.",
    )
    insert_paragraph_before(
        document,
        model_config_paragraph,
        "Seventh, Phase 6 result tables and figures are archived with the "
        "experiment, but package versions remain unpinned and a clean "
        "end-to-end rerun has not been documented. The saved outputs are "
        "internally consistent and reproducible from the station-level CSV, "
        "yet final submission should record the execution environment and "
        "dataset retrieval procedure.",
    )

    # Keep the first practical-interpretation paragraph intact so LibreOffice
    # does not squeeze its heading and opening line together at a page boundary.
    practical_heading = paragraph_starting(
        document, "7.3 Practical meaning of the reported error"
    )
    practical_heading.paragraph_format.keep_with_next = True
    practical_opening = paragraph_starting(
        document, "In the stricter Phase 6 experiment"
    )
    practical_opening.paragraph_format.keep_together = True

    # Avoid leaving only the first line of the introduction at the foot of the
    # abstract page.
    intro_opening = paragraph_starting(
        document, "Rooftop photovoltaic generation is inherently variable"
    )
    intro_opening.paragraph_format.keep_together = True

    # Compact the bibliography slightly so the final citation is not stranded
    # by itself on an otherwise blank page.
    references_heading = paragraph_starting(document, "References")
    references_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph._p is references_heading._p
    )
    for reference in document.paragraphs[references_index + 1 :]:
        if not reference.text.strip():
            continue
        reference.paragraph_format.space_after = Pt(4)
        reference.paragraph_format.line_spacing = 1.1
        for run in reference.runs:
            run.font.size = Pt(10.5)

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
